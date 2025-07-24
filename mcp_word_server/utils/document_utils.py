"""
Document utility functions for Word Document Server.

This module provides utility functions for working with Word documents,
including extracting properties, text, and structure from .docx files.
"""
import os
import functools
from typing import Dict, List, Any, Callable, TypeVar, cast
from docx import Document
from docx.document import Document as DocumentType
from docx.oxml.coreprops import CT_CoreProperties
from docx.section import Section


# Type variable for generic function type
F = TypeVar('F', bound=Callable[..., Any])


def validate_documentpath(param_name: str) -> Callable[[F], F]:
    """
    Decorador parametrizable para validar que el archivo existe
    en el path especificado en el argumento `param_name`.

    Args:
        param_name: nombre del parámetro que contiene el path.

    Returns:
        Función decoradora.
    """
    def decorator(func: F) -> F:
        @functools.wraps(func)
        def wrapper(*args: Any, **kwargs: Any) -> Any:
            # Obtener el valor del parámetro desde kwargs
            path_value = kwargs.get(param_name)

            # Si no está en kwargs, buscar en args por nombre de parámetro
            if path_value is None and args:
                sig = inspect.signature(func)
                param_list = list(sig.parameters.keys())
                if param_name in param_list:
                    index = param_list.index(param_name)
                    if index < len(args):
                        path_value = args[index]

            # Si aún no hay valor, se ejecuta igual sin validar
            if path_value is None:
                return func(*args, **kwargs)

            # Validar que exista el archivo
            if not os.path.exists(path_value):
                return {"error": f"El archivo '{path_value}' no existe."}

            return func(*args, **kwargs)
        return cast(F, wrapper)
    return decorator


@validate_document_path('doc_path')
def get_document_properties(doc_path: str) -> Dict[str, Any]:
    """Get properties of a Word document.
    
    Args:
        doc_path: Path to the Word document.
        
    Returns:
        Dict containing document properties including title, author, subject, etc.
        On error, returns a dict with an 'error' key.
    """
    try:
        doc: DocumentType = Document(doc_path)
        core_props: CT_CoreProperties = doc.core_properties
        sections: List[Section] = doc.sections
        
        word_count: int = sum(
            len(paragraph.text.split()) 
            for paragraph in doc.paragraphs
        )
        
        return {
            "title": core_props.title or "",
            "author": core_props.author or "",
            "subject": core_props.subject or "",
            "keywords": core_props.keywords or "",
            "created": str(core_props.created) if core_props.created else "",
            "modified": str(core_props.modified) if core_props.modified else "",
            "last_modified_by": core_props.last_modified_by or "",
            "revision": core_props.revision or 0,
            "page_count": len(sections),
            "word_count": word_count,
            "paragraph_count": len(doc.paragraphs),
            "table_count": len(doc.tables)
        }
    except Exception as e:
        return {"error": f"Failed to get document properties: {str(e)}"}


@validate_document_path('doc_path')
def extract_document_text(doc_path: str) -> str:
    """Extract all text from a Word document.
    
    Args:
        doc_path: Path to the Word document.
        
    Returns:
        String containing all text content from the document.
        On error, returns an error message string.
    """
    try:
        doc: DocumentType = Document(doc_path)
        text_parts: List[str] = []
        
        # Extract text from paragraphs
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_parts.append(paragraph.text)
        
        # Extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        if paragraph.text.strip():
                            text_parts.append(paragraph.text)
        
        return "\n".join(text_parts)
    except Exception as e:
        return f"Failed to extract text: {str(e)}"


@validate_document_path('doc_path')
def get_document_structure(doc_path: str) -> Dict[str, Any]:
    """Get the structure of a Word document.
    
    Args:
        doc_path: Path to the Word document.
        
    Returns:
        Dict containing document structure including paragraphs and tables.
        On error, returns a dict with an 'error' key.
    """
    try:
        doc: DocumentType = Document(doc_path)
        structure: Dict[str, List[Dict[str, Any]]] = {
            "paragraphs": [],
            "tables": []
        }
        
        # Get paragraphs with preview text
        for para_idx, paragraph in enumerate(doc.paragraphs):
            if not paragraph.text.strip():
                continue
                
            preview_text: str = (
                paragraph.text[:100] + 
                ("..." if len(paragraph.text) > 100 else "")
            )
            
            structure["paragraphs"].append({
                "index": para_idx,
                "text": preview_text,
                "style": paragraph.style.name if paragraph.style else "Normal"
            })
        
        # Get tables with preview data
        for table_idx, table in enumerate(doc.tables):
            table_data: Dict[str, Any] = {
                "index": table_idx,
                "rows": len(table.rows),
                "columns": len(table.columns),
                "preview": []
            }
            
            # Get sample of table data (first 3 rows x first 3 columns)
            max_preview_rows: int = min(3, len(table.rows))
            max_preview_cols: int = min(3, len(table.columns))
            
            for row_idx in range(max_preview_rows):
                row_data: List[str] = []
                for col_idx in range(max_preview_cols):
                    try:
                        cell_text: str = table.cell(row_idx, col_idx).text
                        preview_text: str = (
                            cell_text[:20] + 
                            ("..." if len(cell_text) > 20 else "")
                        )
                        row_data.append(preview_text)
                    except IndexError:
                        row_data.append("N/A")
                
                if row_data:  # Only add non-empty rows
                    table_data["preview"].append(row_data)
            
            structure["tables"].append(table_data)
        
        return structure
    except Exception as e:
        return {"error": f"Failed to get document structure: {str(e)}"}


def find_paragraph_by_text(
    doc: DocumentType, 
    text: str, 
    partial_match: bool = False
) -> List[int]:
    """Find paragraphs containing specific text.
    
    Args:
        doc: Document object to search within.
        text: Text to search for in paragraphs.
        partial_match: If True, matches paragraphs containing the text.
                     If False, matches paragraphs with exact text.
        
    Returns:
        List of paragraph indices (0-based) that match the search criteria.
    """
    if not text or not hasattr(doc, 'paragraphs'):
        return []
    
    matching_paragraphs: List[int] = []
    
    for para_idx, paragraph in enumerate(doc.paragraphs):
        if not paragraph.text:
            continue
            
        if (partial_match and text in paragraph.text) or \
           (not partial_match and paragraph.text == text):
            matching_paragraphs.append(para_idx)
    
    return matching_paragraphs


def find_and_replace_text(
    doc: DocumentType,
    old_text: str,
    new_text: str
) -> int:
    """Find and replace text throughout the document.
    
    Args:
        doc: Document object to search and modify.
        old_text: Text to find in the document.
        new_text: Text to replace the found text with.
        
    Returns:
        int: Number of replacements made.
        
    Raises:
        AttributeError: If the document object is invalid.
        ValueError: If old_text is empty.
    """
    if not old_text:
        raise ValueError("Search text cannot be empty")
    
    if not hasattr(doc, 'paragraphs') or not hasattr(doc, 'tables'):
        raise AttributeError("Invalid document object provided")
    
    replacement_count: int = 0
    
    # Search and replace in paragraphs
    for paragraph in doc.paragraphs:
        if old_text in paragraph.text:
            for run in paragraph.runs:
                if old_text in run.text:
                    run.text = run.text.replace(old_text, new_text)
                    replacement_count += 1
    
    # Search and replace in tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    if old_text in paragraph.text:
                        for run in paragraph.runs:
                            if old_text in run.text:
                                run.text = run.text.replace(old_text, new_text)
                                replacement_count += 1
    
    return replacement_count
