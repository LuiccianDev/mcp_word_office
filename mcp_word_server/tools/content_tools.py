"""Content tools for Word Document Server.

This module provides functions to add and manipulate various types of content
in Word documents, including headings, paragraphs, tables, images, and page breaks.
"""

# modulos standar
import os
from typing import Any, Dict, List, Optional

# modulos terceros
from docx import Document
from docx.document import Document as DocumentType
from docx.shared import Inches, Pt

from mcp_word_server.core.styles import ensure_heading_style
from mcp_word_server.utils.document_utils import find_and_replace_text

# modulos propios
from mcp_word_server.validation.document_validators import (
    check_file_writeable,
    validate_docx_file,
)


@check_file_writeable("filename")
@validate_docx_file("filename")
async def add_heading(filename: str, text: str, level: int = 1) -> str:
    """Add a heading to a Word document.

    Args:
        filename: Path to the Word document as a string.
        text: The text content of the heading.
        level: The heading level (1-9, where 1 is the highest level).

    Returns:
        str: Success message with details of the operation.

    Raises:
        ValueError: If the heading level is invalid.
        IOError: If there's an error processing the document.
    """
    try:
        # Validate and normalize heading level
        heading_level = _validate_heading_level(level)

        doc: DocumentType = Document(filename)
        ensure_heading_style(doc)

        try:
            # Try with style-based approach first
            doc.add_heading(text, level=heading_level)
            doc.save(filename)
            return f"Heading '{text}' (level {heading_level}) added to {filename}"

        except Exception:
            # Fallback to direct formatting if style-based approach fails
            paragraph = doc.add_paragraph(text)
            paragraph.style = doc.styles["Normal"]
            run = paragraph.runs[0]
            run.bold = True

            # Set font size based on heading level
            font_sizes = {1: 16, 2: 14}
            run.font.size = Pt(font_sizes.get(heading_level, 12))

            doc.save(filename)
            return (
                f"Heading '{text}' added to {filename} with direct formatting "
                "(style not available)"
            )

    except Exception as error:
        error_msg = f"Failed to add heading to document: {str(error)}"
        return error_msg


@check_file_writeable("filename")
@validate_docx_file("filename")
async def add_paragraph(filename: str, text: str, style: Optional[str] = None) -> str:
    """Add a paragraph to a Word document.

    Args:
        filename: Path to the Word document as a string.
        text: The text content of the paragraph.
        style: Optional name of the paragraph style to apply.

    Returns:
        str: Success message with details of the operation.

    Raises:
        IOError: If there's an error processing the document.
    """
    try:
        doc: DocumentType = Document(filename)
        paragraph = doc.add_paragraph(text)

        if style:
            try:
                paragraph.style = style
            except KeyError:
                paragraph.style = doc.styles["Normal"]
                doc.save(filename)
                return f"Style '{style}' not found. Paragraph added with default style to {filename}"

        doc.save(filename)
        return f"Paragraph added to {filename}"

    except Exception as error:
        error_msg = f"Failed to add paragraph to document: {str(error)}"
        return error_msg


def _populate_table(
    table: Any,  # docx.table.Table type
    data: List[List[Any]],
    max_rows: int,
    max_cols: int,
) -> None:
    """Populate a table with data.

    Args:
        table: The table to populate.
        data: 2D array of data to fill the table.
        max_rows: Maximum number of rows to populate.
        max_cols: Maximum number of columns to populate.
    """
    for i, row_data in enumerate(data):
        if i >= max_rows:
            break
        for j, cell_text in enumerate(row_data):
            if j >= max_cols:
                break
            table.cell(i, j).text = str(cell_text)


@check_file_writeable("filename")
@validate_docx_file("filename")
async def add_table(
    filename: str, rows: int, cols: int, data: Optional[List[List[Any]]] = None
) -> str:
    """Add a table to a Word document.

    Args:
        filename: Path to the Word document as a string.
        rows: Number of rows in the table.
        cols: Number of columns in the table.
        data: Optional 2D array of data to fill the table.

    Returns:
        str: Success message with details of the operation.

    Raises:
        ValueError: If rows or cols are not positive integers.
        IOError: If there's an error processing the document.
    """
    try:
        # Validate dimensions
        if rows <= 0 or cols <= 0:
            raise ValueError(f"Table dimensions must be positive (got {rows}x{cols})")

        doc: DocumentType = Document(filename)
        table = doc.add_table(rows=rows, cols=cols)

        # Apply table style if available
        try:
            table.style = "Table Grid"
        except KeyError:
            pass  # Use default table style

        # Fill table with data if provided
        if data:
            _populate_table(table, data, rows, cols)

        doc.save(filename)
        return f"Table ({rows}x{cols}) added to {filename}"

    except Exception as error:
        error_msg = f"Failed to add table to document: {str(error)}"

        return error_msg


@check_file_writeable("filename")
@validate_docx_file("filename")
async def add_picture(
    filename: str, image_path: str, width: Optional[float] = None
) -> str:
    """Add an image to a Word document.

    Args:
        filename: Path to the Word document as a string.
        image_path: Path to the image file as a string.
        width: Optional width in inches (proportional scaling).

    Returns:
        str: Success message with details of the operation.

    Raises:
        FileNotFoundError: If the image file doesn't exist.
        ValueError: If width is not a positive number.
        IOError: If there's an error processing the document or image.
    """
    try:
        # Validate image file and get absolute path
        abs_image_path = _validate_image_file(image_path)

        # Validate width if provided
        if width is not None and width <= 0:
            raise ValueError(f"Width must be positive, got {width}")

        doc: DocumentType = Document(filename)

        # Add picture with optional width
        if width is not None:
            doc.add_picture(abs_image_path, width=Inches(width))
        else:
            doc.add_picture(abs_image_path)

        doc.save(filename)
        return f"Picture '{os.path.basename(image_path)}' added to {filename}"

    except Exception as error:
        error_msg = f"Failed to add picture to document: {str(error)}"
        return error_msg


@check_file_writeable("filename")
@validate_docx_file("filename")
async def add_page_break(filename: str) -> str:
    """Add a page break to the document.

    Args:
        filename: Path to the Word document as a string.

    Returns:
        str: Success message confirming the page break was added.

    Raises:
        IOError: If there's an error processing the document.
    """
    try:
        doc: DocumentType = Document(filename)
        doc.add_page_break()
        return f"Page break added to {filename}"
    except Exception as error:
        error_msg = f"Failed to add page break to document: {str(error)}"
        return error_msg


@check_file_writeable("filename")
@validate_docx_file("filename")
async def add_table_of_contents(
    filename: str, title: str = "Table of Contents", max_level: int = 3
) -> str:
    """Add a table of contents to a Word document based on heading styles.

    Args:
        filename: Path to the Word document as a string.
        title: Title for the table of contents.
        max_level: Maximum heading level to include (1-9).

    Returns:
        str: Success message with details of the operation.

    Raises:
        ValueError: If max_level is not between 1 and 9.
        IOError: If there's an error processing the document.
    """
    try:
        # Validate max_level
        if not 1 <= max_level <= 9:
            raise ValueError(f"max_level must be between 1 and 9, got {max_level}")

        doc: DocumentType = Document(filename)

        # Extract headings
        headings = _extract_headings(doc, max_level)

        if not headings:
            return f"No headings found in document {filename}. Table of contents not created."

        # Create a new document for TOC
        toc_doc = Document()

        # Add title if provided
        if title:
            toc_doc.add_heading(title, level=1)

        # Add TOC entries with proper indentation
        for heading in headings:
            indent = "    " * (heading["level"] - 1)
            toc_doc.add_paragraph(f"{indent}{heading['text']}")

        # Add page break after TOC
        toc_doc.add_page_break()

        # Copy original document content
        _copy_document_content(doc, toc_doc)

        # Save the document with TOC
        toc_doc.save(filename)

        return f"Table of contents with {len(headings)} entries added to {filename}"

    except Exception as error:
        error_msg = f"Failed to add table of contents to document: {str(error)}"
        return error_msg


@check_file_writeable("filename")
@validate_docx_file("filename")
async def delete_paragraph(filename: str, paragraph_index: int) -> str:
    """Delete a paragraph from a document.

    Args:
        filename: Path to the Word document as a string.
        paragraph_index: Index of the paragraph to delete (0-based).

    Returns:
        str: Success message with details of the operation.

    Raises:
        IndexError: If paragraph_index is out of range.
        IOError: If there's an error processing the document.
    """
    try:
        doc: DocumentType = Document(filename)
        total_paragraphs = len(doc.paragraphs)

        # Validate paragraph index
        if paragraph_index < 0 or paragraph_index >= total_paragraphs:
            raise IndexError(
                f"Paragraph index {paragraph_index} is out of range. "
                f"Document has {total_paragraphs} paragraphs (0-{total_paragraphs-1})."
            )

        # Remove the paragraph by removing its XML element
        paragraph = doc.paragraphs[paragraph_index]
        paragraph_element = paragraph._p
        paragraph_element.getparent().remove(paragraph_element)

        doc.save(filename)
        return f"Paragraph at index {paragraph_index} deleted successfully."

    except Exception as error:
        error_msg = f"Failed to delete paragraph from document: {str(error)}"
        return error_msg


@check_file_writeable("filename")
@validate_docx_file("filename")
async def search_and_replace(filename: str, find_text: str, replace_text: str) -> str:
    """Search for text and replace all occurrences in a Word document.

    Args:
        filename: Path to the Word document as a string.
        find_text: Text to search for.
        replace_text: Text to replace with.

    Returns:
        str: Success message with details of the operation.

    Raises:
        ValueError: If find_text is empty.
        IOError: If there's an error processing the document.
    """
    try:
        # Validate input
        if not find_text:
            raise ValueError("Search text cannot be empty")

        doc: DocumentType = Document(filename)

        # Perform find and replace
        replacement_count = find_and_replace_text(doc, find_text, replace_text)

        if replacement_count > 0:
            doc.save(filename)
            return (
                f"Replaced {replacement_count} occurrence(s) of "
                f"'{find_text}' with '{replace_text}' in {filename}"
            )

        return f"No occurrences of '{find_text}' found in {filename}"

    except Exception as error:
        error_msg = f"Failed to perform search and replace in document: {str(error)}"
        return error_msg


def _validate_heading_level(level: int) -> int:
    """Validate and normalize heading level.

    Args:
        level: The heading level to validate.

    Returns:
        int: The validated heading level.

    Raises:
        ValueError: If the level is not a valid integer between 1 and 9.
    """
    try:
        level_int = int(level)
        if not 1 <= level_int <= 9:
            raise ValueError(f"Heading level must be between 1 and 9, got {level_int}")
        return level_int
    except (ValueError, TypeError) as error:
        raise ValueError(
            f"Invalid heading level: {level}. Must be an integer between 1 and 9."
        ) from error


def _copy_document_content(source_doc: Any, target_doc: Any) -> None:
    """Copy content from source document to target document.

    Args:
        source_doc: Source document to copy from.
        target_doc: Target document to copy to.
    """
    # Copy paragraphs
    for paragraph in source_doc.paragraphs:
        new_paragraph = target_doc.add_paragraph(paragraph.text)
        if paragraph.style:
            try:
                new_paragraph.style = paragraph.style.name
            except KeyError:
                pass  # Use default style

    # Copy tables
    for table in source_doc.tables:
        new_table = target_doc.add_table(rows=len(table.rows), cols=len(table.columns))
        for i, row in enumerate(table.rows):
            for j, cell in enumerate(row.cells):
                new_table.cell(i, j).text = "\n".join(p.text for p in cell.paragraphs)


def _extract_headings(
    doc: Any, max_level: int  # docx.document.Document type
) -> List[Dict[str, Any]]:
    """Extract headings from a document.

    Args:
        doc: The document to extract headings from.
        max_level: Maximum heading level to include.

    Returns:
        List of dictionaries containing heading information.
    """
    headings = []
    for i, paragraph in enumerate(doc.paragraphs):
        if paragraph.style and paragraph.style.name.startswith("Heading "):
            try:
                level = int(paragraph.style.name.split(" ")[1])
                if level <= max_level:
                    headings.append(
                        {"level": level, "text": paragraph.text, "position": i}
                    )
            except (ValueError, IndexError):
                continue
    return headings


def _validate_image_file(image_path: str) -> str:
    """Validate an image file and return its absolute path.

    Args:
        image_path: Path to the image file to validate.

    Returns:
        str: Absolute path to the image file.

    Raises:
        FileNotFoundError: If the image file doesn't exist.
        IOError: If there's an error accessing the image file.
    """
    abs_image_path = os.path.abspath(image_path)

    if not os.path.exists(abs_image_path):
        raise FileNotFoundError(f"Image file not found: {abs_image_path}")

    # Check file size
    try:
        image_size_kb = os.path.getsize(abs_image_path) / 1024
        if image_size_kb <= 0:
            raise IOError(f"Image file is empty: {abs_image_path}")
        return abs_image_path

    except OSError as error:
        raise IOError(f"Error accessing image file: {error}") from error


def _validate_image_file(image_path: str) -> str:
    """Validate an image file and return its absolute path.

    Args:
        image_path: Path to the image file to validate.

    Returns:
        str: Absolute path to the image file.

    Raises:
        FileNotFoundError: If the image file doesn't exist.
        IOError: If there's an error accessing the image file.
    """
    abs_image_path = os.path.abspath(image_path)

    if not os.path.exists(abs_image_path):
        raise FileNotFoundError(f"Image file not found: {abs_image_path}")

    # Check file size
    try:
        image_size_kb = os.path.getsize(abs_image_path) / 1024
        if image_size_kb <= 0:
            raise IOError(f"Image file is empty: {abs_image_path}")
        return abs_image_path

    except OSError as error:
        raise IOError(f"Error accessing image file: {error}") from error
