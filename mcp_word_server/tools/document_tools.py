"""Document creation and manipulation tools for Word Document Server.

This module provides functions for creating, reading, and manipulating
Word documents with security checks for file system access.
"""

# modulos estandar
import json
import os
from typing import Any, Dict, List, Optional, Tuple

# modulos de terceros
from docx import Document
from docx.document import Document as DocumentType

from mcp_word_server.core.styles import ensure_heading_style, ensure_table_style

# modulos propios
from mcp_word_server.core.tables import copy_table
from mcp_word_server.utils.document_utils import (
    create_document_copy,
    ensure_docx_extension,
    extract_document_text,
    get_document_properties,
    get_document_structure,
)
from mcp_word_server.validation.document_validators import (
    check_file_writeable,
    validate_docx_file,
)


def _get_allowed_directories() -> List[str]:
    """Get the list of allowed directories from environment variables.

    Returns:
        List of absolute paths to directories where documents can be created/accessed.
        Defaults to ['./documents'] if MCP_ALLOWED_DIRECTORIES is not set.
    """

    allowed_dirs_str = os.environ.get("MCP_ALLOWED_DIRECTORIES", "./documents")

    allowed_dirs = [dir.strip() for dir in allowed_dirs_str.split(",")]

    return [os.path.abspath(dir) for dir in allowed_dirs]


def _is_path_in_allowed_directories(file_path: str) -> Tuple[bool, Optional[str]]:
    """Check if the given file path is within allowed directories.

    Args:
        file_path: The file path to validate.

    Returns:
        Tuple of (is_allowed, error_message) where is_allowed is a boolean
        indicating if the path is allowed, and error_message provides details
        if the path is not allowed.
    """
    allowed_dirs = _get_allowed_directories()
    abs_path = os.path.abspath(file_path)

    for allowed_dir in allowed_dirs:
        try:
            if os.path.commonpath([allowed_dir, abs_path]) == allowed_dir:
                return True, None
        except ValueError:
            continue

    error_msg = (
        f"Path '{file_path}' is not in allowed directories: "
        f"{', '.join(allowed_dirs)}"
    )
    return False, error_msg


async def create_document(
    filename: str, title: Optional[str] = None, author: Optional[str] = None
) -> str:
    """Create a new Word document with optional metadata.

    Args:
        filename: Name of the document to create (with or without .docx extension).
        title: Optional title for the document metadata.
        author: Optional author for the document metadata.

    Returns:
        str: Success or error message indicating the result of the operation.
    """

    is_allowed, error_message = _is_path_in_allowed_directories(filename)
    if not is_allowed:
        return f"Cannot create document: {error_message}"

    directory = os.path.dirname(filename)
    if directory and not os.path.exists(directory):
        try:
            os.makedirs(directory, exist_ok=True)
        except OSError as e:
            return f"Cannot create directory '{directory}': {str(e)}"

    try:
        doc: DocumentType = Document()

        # Set properties if provided
        if title:
            doc.core_properties.title = title
        if author:
            doc.core_properties.author = author

        # Ensure necessary styles exist
        ensure_heading_style(doc)
        ensure_table_style(doc)

        # Save the document
        doc.save(filename)

        return f"Document {filename} created successfully"
    except Exception as e:
        return f"Failed to create document: {str(e)}"

    #! Function removed
    # async def create_document(filename: str, title: Optional[str] = None, author: Optional[str] = None) -> str:
    """Create a new Word document with optional metadata.
    
    Args:
        filename: Name of the document to create (with or without .docx extension)
        title: Optional title for the document metadata
        author: Optional author for the document metadata
    """
    """ filename = ensure_docx_extension(filename)
    
    # Check if file is writeable
    is_writeable, error_message = check_file_writeable(filename)
    if not is_writeable:
        return f"Cannot create document: {error_message}"
    
    try:
        doc = Document()
        
        # Set properties if provided
        if title:
            doc.core_properties.title = title
        if author:
            doc.core_properties.author = author
        
        # Ensure necessary styles exist
        ensure_heading_style(doc)
        ensure_table_style(doc)
        
        # Save the document
        doc.save(filename)
        
        return f"Document {filename} created successfully"
    except Exception as e:
        return f"Failed to create document: {str(e)}" """


@validate_docx_file("filename")
async def get_document_info(filename: str) -> str:
    """Get information about a Word document.

    Args:
        filename: Path to the Word document
    """
    try:
        properties = get_document_properties(filename)
        return json.dumps(properties, indent=2)
    except Exception as e:
        return f"Failed to get document info: {str(e)}"


@validate_docx_file("filename")
async def get_document_text(filename: str) -> str:
    """Extract all text from a Word document.

    Args:
        filename: Path to the Word document
    """
    return extract_document_text(filename)


@validate_docx_file("filename")
async def get_document_outline(filename: str) -> str:
    """Get the structure of a Word document.

    Args:
        filename: Path to the Word document
    """
    structure = get_document_structure(filename)
    return json.dumps(structure, indent=2)


async def list_available_documents(directory: Optional[str] = None) -> Dict[str, Any]:
    """
    List all .docx files in the specified or allowed directories.

    Args:
        directory: Optional path to a directory. If not provided, uses all allowed.

    Returns:
        Dictionary with status, message, and list of found documents.
    """
    try:
        # Si no se especifica directorio, usamos todos los permitidos
        if directory is None:
            search_directories = _get_allowed_directories()
            if not search_directories:
                return {
                    "status": "error",
                    "message": "No accessible directories found in MCP configuration.",
                    "allowed_directories": _get_allowed_directories(),
                    "documents": [],
                }
        else:
            abs_dir = os.path.abspath(directory)
            allowed_dirs = _get_allowed_directories()

            # Validación segura usando commonpath
            is_allowed = any(
                os.path.commonpath([abs_dir, allowed]) == allowed
                for allowed in allowed_dirs
            )
            if not is_allowed:
                return {
                    "status": "error",
                    "message": f"Directory '{directory}' is not in allowed directories: {', '.join(allowed_dirs)}",
                    "allowed_directories": allowed_dirs,
                    "documents": [],
                }
            search_directories = [abs_dir]

        all_documents = []
        total_found = 0

        for search_dir in search_directories:
            if not os.path.exists(search_dir):
                continue

            docx_files = [
                f
                for f in os.listdir(search_dir)
                if f.lower().endswith(".docx")
                and not f.startswith("~$")  # Excluir archivos temporales
            ]

            for file in sorted(docx_files):
                file_path = os.path.join(search_dir, file)
                size_kb = os.path.getsize(file_path) / 1024
                all_documents.append(
                    {
                        "name": file,
                        "path": os.path.abspath(file_path),
                        "size_kb": round(size_kb, 2),
                        "source_directory": search_dir,
                    }
                )

            total_found += len(docx_files)

        return {
            "status": "ok",
            "message": f"Found {total_found} Word document(s)"
            + (
                f" across {len(search_directories)} directories"
                if len(search_directories) > 1
                else ""
            )
            + ".",
            "directories_searched": search_directories,
            "total": total_found,
            "documents": all_documents,
        }

    except Exception as e:
        return {
            "status": "error",
            "message": f"Failed to list documents: {str(e)}",
            "documents": [],
        }


@validate_docx_file("source_filename")
async def copy_document(
    source_filename: str, destination_filename: Optional[str] = None
) -> str:
    """Create a copy of a Word document.

    Args:
        source_filename: Path to the source document
        destination_filename: Optional path for the copy. If not provided, a default name will be generated.
    """
    source_filename = ensure_docx_extension(source_filename)

    if destination_filename:
        destination_filename = ensure_docx_extension(destination_filename)

    success, message, _ = create_document_copy(source_filename, destination_filename)
    if success:
        return message
    return f"Failed to copy document: {message}"


@check_file_writeable("target_filename")
@validate_docx_file("target_filename")
async def merge_documents(
    target_filename: str, source_filenames: List[str], add_page_breaks: bool = True
) -> str:
    """Merge multiple Word documents into a single document.

    Args:
        target_filename: Path to the target document (will be created or overwritten)
        source_filenames: List of paths to source documents to merge
        add_page_breaks: If True, add page breaks between documents
    """
    # Validate all source documents exist
    missing_files = []
    for filename in source_filenames:
        doc_filename = ensure_docx_extension(filename)
        if not os.path.exists(doc_filename):
            missing_files.append(doc_filename)

    if missing_files:
        return f"Cannot merge documents. The following source files do not exist: {', '.join(missing_files)}"

    try:
        # Create a new document for the merged result
        target_doc: DocumentType = Document()

        # Process each source document
        for i, filename in enumerate(source_filenames):
            doc_filename = ensure_docx_extension(filename)
            source_doc: DocumentType = Document(doc_filename)

            # Add page break between documents (except before the first one)
            if add_page_breaks and i > 0:
                target_doc.add_page_break()

            # Copy all paragraphs
            for paragraph in source_doc.paragraphs:
                # Create a new paragraph with the same text and style
                new_paragraph = target_doc.add_paragraph(paragraph.text)
                new_paragraph.style = target_doc.styles["Normal"]  # Default style

                # Try to match the style if possible
                try:
                    if paragraph.style and paragraph.style.name in target_doc.styles:
                        new_paragraph.style = target_doc.styles[paragraph.style.name]
                except:
                    pass

                # Copy run formatting
                for i, run in enumerate(paragraph.runs):
                    if i < len(new_paragraph.runs):
                        new_run = new_paragraph.runs[i]
                        # Copy basic formatting
                        new_run.bold = run.bold
                        new_run.italic = run.italic
                        new_run.underline = run.underline
                        # Font size if specified
                        if run.font.size:
                            new_run.font.size = run.font.size

            # Copy all tables
            for table in source_doc.tables:
                copy_table(table, target_doc)

        # Save the merged document
        target_doc.save(target_filename)
        return f"Successfully merged {len(source_filenames)} documents into {target_filename}"
    except Exception as e:
        return f"Failed to merge documents: {str(e)}"
