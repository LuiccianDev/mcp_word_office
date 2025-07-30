"""
Protection tools for Word Document Server.

These tools handle document protection features such as
password protection, restricted editing, and digital signatures.
"""

# modulos estandar
import datetime
import hashlib
import io
import os
from typing import Any, List, Optional

import msoffcrypto

# modulos de terceros
from docx import Document
from docx.document import Document as DocumentType

from word_mcp.core.protection import (
    add_protection_info,
    create_signature_info,
    verify_document_protection,
)

# modulos propios
from word_mcp.validation.document_validators import (
    check_file_writeable,
    validate_docx_file,
)


@validate_docx_file("filename")
@check_file_writeable("filename")
async def protect_document(filename: str, password: str) -> dict[str, Any]:
    """Add password protection to a Word document.

    Args:
        filename: Path to the Word document
        password: Password to protect the document with
    """
    try:
        # Read the original file content
        with open(filename, "rb") as infile:
            original_data = infile.read()

        # Create an msoffcrypto file object from the original data
        file = msoffcrypto.OfficeFile(io.BytesIO(original_data))
        file.load_key(password=password)  # Set the password for encryption

        # Encrypt the data into an in-memory buffer
        encrypted_data_io = io.BytesIO()

        file.encrypt(password=password, outfile=encrypted_data_io)

        # Overwrite the original file with the encrypted data
        with open(filename, "wb") as outfile:
            outfile.write(encrypted_data_io.getvalue())

        base_path, _ = os.path.splitext(filename)
        metadata_path = f"{base_path}.protection"
        if os.path.exists(metadata_path):
            os.remove(metadata_path)

        return {"status": "success", "message": f"Document {filename} encrypted successfully with password."}

    except Exception as e:
        # Attempt to restore original file content on failure
        try:
            if "original_data" in locals():
                with open(filename, "wb") as outfile:
                    outfile.write(original_data)
                return {"status": "error", "error": f"Failed to encrypt document {filename}: {str(e)}. Original file restored."}
            else:
                return {"status": "error", "error": f"Failed to encrypt document {filename}: {str(e)}. Could not restore original file."}
        except Exception as restore_e:
            return {"status": "error", "error": f"Failed to encrypt document {filename}: {str(e)}. Also failed to restore original file: {str(restore_e)}"}


@validate_docx_file("filename")
@check_file_writeable("filename")
async def add_restricted_editing(
    filename: str, password: str, editable_sections: List[str]
) -> dict[str, Any]:
    """Add restricted editing to a Word document, allowing editing only in specified sections.

    Args:
        filename: Path to the Word document
        password: Password to protect the document with
        editable_sections: List of section names that can be edited
    """
    try:
        # Hash the password for security
        password_hash = hashlib.sha256(password.encode()).hexdigest()

        # Add protection info to metadata
        success = add_protection_info(
            filename,
            protection_type="restricted",
            password_hash=password_hash,
            sections=editable_sections,
        )

        if not editable_sections:
            return {"status": "error", "error": "No editable sections specified. Document will be fully protected."}

        if success:
            return {"status": "success", "message": f"Document {filename} protected with restricted editing. Editable sections: {', '.join(editable_sections)}"}
        else:
            return {"status": "error", "error": f"Failed to protect document {filename} with restricted editing"}
    except Exception as e:
        return {"status": "error", "error": f"Failed to add restricted editing: {str(e)}"}


@validate_docx_file("filename")
@check_file_writeable("filename")
async def add_digital_signature(
    filename: str, signer_name: str, reason: Optional[str] = None
) -> dict[str, Any]:
    """Add a digital signature to a Word document.

    Args:
        filename: Path to the Word document
        signer_name: Name of the person signing the document
        reason: Optional reason for signing
    """
    try:
        doc: DocumentType = Document(filename)

        # Create signature info
        signature_info = create_signature_info(doc, signer_name, reason)

        # Add protection info to metadata
        success = add_protection_info(
            filename,
            protection_type="signature",
            password_hash="",  # No password for signature-only
            signature_info=signature_info,
        )

        if success:
            # Add a visible signature block to the document
            doc.add_paragraph("").add_run()  # Add empty paragraph for spacing
            signature_para = doc.add_paragraph()
            signature_para.add_run(f"Digitally signed by: {signer_name}").bold = True
            if reason:
                signature_para.add_run(f"\nReason: {reason}")
            signature_para.add_run(
                f"\nDate: {datetime.datetime.now().strftime('%Y-%m-%d %H:%M:%S')}"
            )
            signature_para.add_run(
                f"\nSignature ID: {signature_info['content_hash'][:8]}"
            )

            # Save the document with the visible signature
            doc.save(filename)

            return {"status": "success", "message": f"Digital signature added to document {filename}"}
        else:
            return {"status": "error", "error": f"Failed to add digital signature to document {filename}"}
    except Exception as e:
        return {"status": "error", "error": f"Failed to add digital signature: {str(e)}"}


@validate_docx_file("filename")
async def verify_document(filename: str, password: Optional[str] = None) -> dict[str, Any]:
    """Verify document protection and/or digital signature.

    Args:
        filename: Path to the Word document
        password: Optional password to verify
    """

    try:
        # Verify document protection
        is_verified, message = verify_document_protection(filename, password)

        if not is_verified and password:
            return {"status": "error", "error": f"Document verification failed: {message}"}

        # If document has a digital signature, verify content integrity
        base_path, _ = os.path.splitext(filename)
        metadata_path = f"{base_path}.protection"

        if os.path.exists(metadata_path):
            try:
                import json

                with open(metadata_path, "r") as f:
                    protection_data = json.load(f)

                if protection_data.get("type") == "signature":
                    # Get the original content hash
                    signature_info = protection_data.get("signature", {})
                    original_hash = signature_info.get("content_hash")

                    if original_hash:
                        # Calculate current content hash
                        doc = Document(filename)
                        text_content = "\n".join([p.text for p in doc.paragraphs])
                        current_hash = hashlib.sha256(text_content.encode()).hexdigest()

                        # Compare hashes
                        if current_hash != original_hash:
                            return {"status": "error", "error": f"Document has been modified since it was signed by {signature_info.get('signer')}"}
                        else:
                            return {"status": "success", "message": f"Document signature is valid. Signed by {signature_info.get('signer')} on {signature_info.get('timestamp')}"}
            except Exception as e:
                return {"status": "error", "error": f"Error verifying signature: {str(e)}"}

        return {"status": "success", "message": message}
    except Exception as e:
        return {"status": "error", "error": f"Failed to verify document: {str(e)}"}


@validate_docx_file("filename")
@check_file_writeable("filename")
async def unprotect_document(filename: str, password: str) -> dict[str, Any]:
    """Remove password protection from a Word document.

    Args:
        filename: Path to the Word document
        password: Password that was used to protect the document
    """
    try:
        # Read the encrypted file content
        with open(filename, "rb") as infile:
            encrypted_data = infile.read()

        # Create an msoffcrypto file object from the encrypted data
        file = msoffcrypto.OfficeFile(io.BytesIO(encrypted_data))
        file.load_key(password=password)  # Set the password for decryption

        # Decrypt the data into an in-memory buffer
        decrypted_data_io = io.BytesIO()
        file.decrypt(
            outfile=decrypted_data_io
        )  # Pass the buffer as the 'outfile' argument

        # Overwrite the original file with the decrypted data
        with open(filename, "wb") as outfile:
            outfile.write(decrypted_data_io.getvalue())

        return {"status": "success", "message": f"Document {filename} decrypted successfully."}

    except msoffcrypto.exceptions.InvalidKeyError:
        return {"status": "error", "error": f"Failed to decrypt document {filename}: Incorrect password."}
    except msoffcrypto.exceptions.InvalidFormatError:
        return {"status": "error", "error": f"Failed to decrypt document {filename}: File is not encrypted or is not a supported Office format."}
    except Exception as e:
        # Attempt to restore encrypted file content on failure
        try:
            if "encrypted_data" in locals():
                with open(filename, "wb") as outfile:
                    outfile.write(encrypted_data)
                return {"status": "error", "error": f"Failed to decrypt document {filename}: {str(e)}. Encrypted file restored."}
            else:
                return {"status": "error", "error": f"Failed to decrypt document {filename}: {str(e)}. Could not restore encrypted file."}
        except Exception as restore_e:
            return {"status": "error", "error": f"Failed to decrypt document {filename}: {str(e)}. Also failed to restore encrypted file: {str(restore_e)}"}
