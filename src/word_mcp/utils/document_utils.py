"""
Document utility functions for Word Document Server.

This module provides utility functions for working with Word documents,
including extracting properties, text, and structure from .docx files.
"""

import os
import shutil
from typing import Any, Dict, List, Optional, Tuple

from docx import Document
from docx.document import Document as DocumentType
from docx.oxml.coreprops import CT_CoreProperties
from docx.section import Section

from word_mcp.validation.document_validators import validate_docx_file


@validate_docx_file("filename")
def get_document_properties(filename: str) -> Dict[str, Any]:
    """Get properties of a Word document.

    Args:
        filename: Path to the Word document.

    Returns:
        Dict containing document properties including title, author, subject, etc.
        On error, returns a dict with an 'error' key.
    """
    try:
        doc: DocumentType = Document(filename)
        core_props: CT_CoreProperties = doc.core_properties
        sections: List[Section] = doc.sections

        word_count: int = sum(
            len(paragraph.text.split()) for paragraph in doc.paragraphs
        )

        return {
            "title": core_props.title or "",
            "author": core_props.author or "",
            "subject": core_props.subject or "",
            "keywords": core_props.keywords or "",
            "created": str(core_props.created) if core_props.created else "",
            "modified": str(core_props.modified) if core_props.modified else "",
            "last_modified_by": core_props.last_modified_by or "",
            "revision": core_props.revision or 0,
            "page_count": len(sections),
            "word_count": word_count,
            "paragraph_count": len(doc.paragraphs),
            "table_count": len(doc.tables),
        }
    except Exception as e:
        return {"error": f"Failed to get document properties: {str(e)}"}


@validate_docx_file("filename")
def extract_document_text(filename: str) -> str:
    """Extract all text from a Word document.

    Args:
        filename: Path to the Word document.

    Returns:
        String containing all text content from the document.
        On error, returns an error message string.
    """
    try:
        doc: DocumentType = Document(filename)
        text_parts: List[str] = []

        # Extract text from paragraphs
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_parts.append(paragraph.text)

        # Extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        if paragraph.text.strip():
                            text_parts.append(paragraph.text)

        return "\n".join(text_parts)
    except Exception as e:
        return f"Failed to extract text: {str(e)}"


@validate_docx_file("filename")
def get_document_structure(filename: str) -> Dict[str, Any]:
    """Get the structure of a Word document.

    Args:
        filename: Path to the Word document.

    Returns:
        Dict containing document structure including paragraphs and tables.
        On error, returns a dict with an 'error' key.
    """
    try:
        doc: DocumentType = Document(filename)
        structure: Dict[str, List[Dict[str, Any]]] = {"paragraphs": [], "tables": []}

        # Get paragraphs with preview text
        for para_idx, paragraph in enumerate(doc.paragraphs):
            if not paragraph.text.strip():
                continue

            preview_text: str = paragraph.text[:100] + (
                "..." if len(paragraph.text) > 100 else ""
            )

            structure["paragraphs"].append(
                {
                    "index": para_idx,
                    "text": preview_text,
                    "style": paragraph.style.name if paragraph.style else "Normal",
                }
            )

        # Get tables with preview data
        for table_idx, table in enumerate(doc.tables):
            table_data: Dict[str, Any] = {
                "index": table_idx,
                "rows": len(table.rows),
                "columns": len(table.columns),
                "preview": [],
            }

            # Get sample of table data (first 3 rows x first 3 columns)
            max_preview_rows: int = min(3, len(table.rows))
            max_preview_cols: int = min(3, len(table.columns))

            for row_idx in range(max_preview_rows):
                row_data: List[str] = []
                for col_idx in range(max_preview_cols):
                    try:
                        cell_text: str = table.cell(row_idx, col_idx).text
                        preview_text: str = cell_text[:20] + (
                            "..." if len(cell_text) > 20 else ""
                        )
                        row_data.append(preview_text)
                    except IndexError:
                        row_data.append("N/A")

                if row_data:  # Only add non-empty rows
                    table_data["preview"].append(row_data)

            structure["tables"].append(table_data)

        return structure
    except Exception as e:
        return {"status": "error", "message": f"Failed to get document structure: {str(e)}"}


def find_paragraph_by_text(
    doc: DocumentType, text: str, partial_match: bool = False
) -> List[int]:
    """Find paragraphs containing specific text.

    Args:
        doc: Document object to search within.
        text: Text to search for in paragraphs.
        partial_match:  If True, matches paragraphs containing the text.
                        If False, matches paragraphs with exact text.

    Returns:
        List of paragraph indices (0-based) that match the search criteria.
    """
    if not text or not hasattr(doc, "paragraphs"):
        return []

    matching_paragraphs: List[int] = []

    for para_idx, paragraph in enumerate(doc.paragraphs):
        if not paragraph.text:
            continue

        if (partial_match and text in paragraph.text) or (
            not partial_match and paragraph.text == text
        ):
            matching_paragraphs.append(para_idx)

    return matching_paragraphs


def find_and_replace_text(doc: DocumentType, old_text: str, new_text: str) -> int:
    """Find and replace text throughout the document.

    Args:
        doc: Document object to search and modify.
        old_text: Text to find in the document.
        new_text: Text to replace the found text with.
    Returns:
        int: Number of replacements made.
    Raises:
        AttributeError: If the document object is invalid.
        ValueError: If old_text is empty.
    """
    if not old_text:
        raise ValueError("Search text cannot be empty")

    if not hasattr(doc, "paragraphs") or not hasattr(doc, "tables"):
        raise AttributeError("Invalid document object provided")

    replacement_count: int = 0

    # Search and replace in paragraphs
    for paragraph in doc.paragraphs:
        if old_text in paragraph.text:
            for run in paragraph.runs:
                if old_text in run.text:
                    run.text = run.text.replace(old_text, new_text)
                    replacement_count += 1

    # Search and replace in tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    if old_text in paragraph.text:
                        for run in paragraph.runs:
                            if old_text in run.text:
                                run.text = run.text.replace(old_text, new_text)
                                replacement_count += 1

    return replacement_count


def create_document_copy(
    source_path: str, dest_path: Optional[str] = None
) -> Tuple[bool, str, Optional[str]]:
    """Create a copy of a document.

    If `dest_path` is not provided, a new filename is generated by appending
    '_copy' to the original filename.

    Args:
        source_path: The path to the source document.
        dest_path: Optional path for the new document.

    Returns:
        A tuple containing a boolean for success, a status message, and the
        path to the new file if successful.
    """
    if not os.path.exists(source_path):
        return False, f"Source document '{source_path}' does not exist.", None

    if not dest_path:
        base, ext = os.path.splitext(source_path)
        dest_path = f"{base}_copy{ext}"

    try:
        shutil.copy2(source_path, dest_path)
        return True, f"Document successfully copied to '{dest_path}'.", dest_path
    except (shutil.Error, IOError) as e:
        return False, f"Failed to copy document: {e}", None


#! Observe that this function is not used in the code, it is left for future use
def ensure_docx_extension(filename: str) -> str:
    """Ensures a filename has a .docx extension.

    If the filename does not already end with '.docx', the extension is appended.
    This is case-sensitive.

    Args:
        filename: The filename to check and modify if necessary.

    Returns:
        The filename with the .docx extension.
    """
    if not filename.lower().endswith(".docx"):
        return filename + ".docx"
    return filename
