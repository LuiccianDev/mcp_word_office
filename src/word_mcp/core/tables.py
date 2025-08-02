"""
Table-related operations for Word Document Server.

This module provides functionality for working with tables in Word documents,
including formatting borders, applying styles, and copying tables between documents.
"""

import re
from typing import Any, Dict, List, Optional, Union

from docx.document import Document
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
from docx.oxml.shared import OxmlElement, qn
from docx.table import Table, _Cell, _Row

# Type aliases for better code readability
CellBorderStyle = Dict[str, Union[str, bool, int]]
TableShading = List[List[Optional[str]]]


def set_cell_border(cell: _Cell, **kwargs: Union[str, bool, int]) -> None:
    """Set cell border properties with validation and type safety.

    Args:
        cell: The table cell to modify.
        **kwargs: Border properties including:
            - top/bottom/left/right: bool - Whether to apply the border
            - val: str - Border style (e.g., 'single', 'double', 'dashed')
            - sz: Union[str, int] - Border width in eighths of a point
            - space: Union[str, int] - Border spacing in points
            - color: str - Border color as hex RGB (e.g., 'FF0000' for red)

    Example:
        set_cell_border(
            cell,
            top=True,
            bottom=True,
            left=True,
            right=True,
            val='single',
            sz='4',
            space='0',
            color='000000'
        )
    """
    tc = cell._tc
    tc_properties = tc.get_or_add_tcPr()

    # Default border properties
    border_props = {
        "val": str(kwargs.get("val", "single")),
        "sz": str(kwargs.get("sz", "4")),
        "space": str(kwargs.get("space", "0")),
        "color": str(kwargs.get("color", "auto")),
    }

    # Get or create borders container
    tc_borders = _get_or_create_tc_borders(tc_properties)

    # Apply borders to specified sides
    for side in ["top", "bottom", "left", "right"]:
        if kwargs.get(side, False):
            border_tag = f"w:{side}"
            border_element = OxmlElement(border_tag)

            # Set border properties
            for attr, value in border_props.items():
                border_element.set(qn(f"w:{attr}"), value)

            # Add or update the border element
            existing_border = tc_borders.find(f"{{*}}{border_tag}")
            if existing_border is not None:
                tc_borders.remove(existing_border)
            tc_borders.append(border_element)


def apply_table_style(
    table: Table,
    has_header_row: bool = False,
    border_style: Optional[str] = None,
    shading: Optional[TableShading] = None,
) -> bool:
    """Apply consistent formatting to a table.

    Args:
        table: The table to format.
        has_header_row: If True, formats the first row as a header.
        border_style: Style for borders ('none', 'single', 'double', 'thick').
        shading: Optional 2D list of cell background colors (by row and column).

    Returns:
        bool: True if formatting was applied successfully, False otherwise.
    """
    try:
        # Format header row if requested
        if has_header_row and table.rows:
            _format_header_row(table.rows[0])

        # Apply borders if specified
        if border_style is not None:
            _apply_border_style(table, border_style)

        # Apply cell shading if specified
        if shading is not None:
            _apply_shading(table, shading)

        return True
    except Exception:
        # Log the error in a real application
        return False


def copy_table(source_table: Table, target_doc: Document) -> Table:
    """Create a deep copy of a table in another document.

    Args:
        source_table: The table to copy.
        target_doc: The document to copy the table to.

    Returns:
        Table: The newly created table in the target document.

    Raises:
        ValueError: If the source table has no rows or columns.
    """
    if not source_table.rows or not source_table.columns:
        raise ValueError("Source table has no rows or columns")

    # Create a new table with matching dimensions
    new_table: Table = target_doc.add_table(
        rows=len(source_table.rows), cols=len(source_table.columns)
    )

    # Apply the same style if possible
    _apply_table_style(source_table, new_table)

    # Copy cell contents and formatting
    for i, row in enumerate(source_table.rows):
        for j, cell in enumerate(row.cells):
            if i < len(new_table.rows) and j < len(new_table.rows[i].cells):
                _copy_table_cell_contents(cell, new_table.cell(i, j))

    return new_table


def _format_header_row(header_row: _Row) -> None:
    """Format a table header row by making all text bold.

    Args:
        header_row: The table row to format as a header.
    """
    for cell in header_row.cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True


def _apply_border_style(table: Table, border_style: str) -> None:
    """Apply consistent border style to all cells in a table.

    Args:
        table: The table to apply borders to.
        border_style: The style of border to apply.
    """
    border_map = {
        "none": "nil",
        "single": "single",
        "double": "double",
        "thick": "thick",
    }

    border_val = border_map.get(border_style.lower(), "single")

    for row in table.rows:
        for cell in row.cells:
            set_cell_border(
                cell=cell,
                top=True,
                bottom=True,
                left=True,
                right=True,
                val=border_val,
                color="000000",
            )


def _apply_shading(table: Table, shading: TableShading) -> None:
    """Apply cell-by-cell shading to a table.

    Args:
        table: The table to apply shading to.
        shading: 2D list of color values (by row and column).
    """
    HEX_COLOR_RE = re.compile(r"^[0-9A-Fa-f]{6}$")

    for i, row_colors in enumerate(shading):
        if i >= len(table.rows):
            break

        for j, color in enumerate(row_colors):
            if j >= len(table.rows[i].cells) or not color:
                continue

            if not HEX_COLOR_RE.match(color):
                continue

            try:
                cell = table.rows[i].cells[j]
                shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}"/>')
                cell._tc.get_or_add_tcPr().append(shading_elm)
            except (ValueError, AttributeError):
                # Skip invalid color formats or malformed XML
                continue


def _apply_table_style(source_table: Table, target_table: Table) -> None:
    """Apply the style from source table to target table.

    Args:
        source_table: The source table with the desired style.
        target_table: The target table to apply the style to.
    """
    try:
        if source_table.style:
            target_table.style = source_table.style
    except (ValueError, KeyError):
        # Fall back to default grid style if source style is not available
        try:
            target_table.style = "Table Grid"
        except (ValueError, KeyError):
            # If Table Grid is not available, use default style
            pass


def _copy_table_cell_contents(source_cell: _Cell, target_cell: _Cell) -> None:
    """Copy text content from source cell to target cell.

    Args:
        source_cell: The cell to copy from.
        target_cell: The cell to copy to.
    """
    # Clear any existing content in the target cell
    target_cell.text = ""

    # Copy paragraph by paragraph to preserve formatting
    for source_paragraph in source_cell.paragraphs:
        if not source_paragraph.text.strip():
            continue

        # Create a new paragraph in the target cell
        target_paragraph = target_cell.add_paragraph()

        # Copy paragraph style and alignment
        target_paragraph.style = source_paragraph.style
        target_paragraph.alignment = source_paragraph.alignment
        # Copy text and formatting from each run
        for run in source_paragraph.runs:
            new_run = target_paragraph.add_run(run.text)
            # Copy run properties
            new_run.bold = run.bold
            new_run.italic = run.italic
            new_run.underline = run.underline
            if run.font:
                new_run.font.name = run.font.name
                new_run.font.size = run.font.size
                new_run.font.color.rgb = run.font.color.rgb


def _get_or_create_tc_borders(tc_properties: Any) -> Any:
    """Get or create the tcBorders element in table cell properties.

    Args:
        tc_properties: The table cell properties element.

    Returns:
        The tcBorders element (existing or newly created).
    """
    borders = tc_properties.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_properties.append(borders)
    return borders
