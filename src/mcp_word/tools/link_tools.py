"""FastMCP tools for hyperlinks and bookmarks manipulation."""

from typing import Any

from docx import Document

from mcp_word.core import links
from mcp_word.validation.document_validators import validate_docx_write


@validate_docx_write("file_path")
async def add_hyperlink(
    file_path: str, paragraph_index: int, text: str, url: str = "", bookmark: str = ""
) -> dict[str, Any]:
    """Add a hyperlink to a specific paragraph.

    You must provide either an external `url` OR an internal `bookmark` name, but not both.

    Args:
        file_path: Path to the .docx file
        paragraph_index: The index of the paragraph to append the hyperlink to (0-based)
        text: The visible text of the hyperlink
        url: Optional external URL (e.g., 'https://example.com')
        bookmark: Optional internal bookmark name to link to

    Returns:
        A dictionary containing the status and a message or error details
    """
    if bool(url) == bool(bookmark):
        return {"status": "error", "message": "Must provide exactly one of 'url' or 'bookmark'."}

    try:
        doc = Document(file_path)
        
        paragraphs = doc.paragraphs
        if paragraph_index < 0 or paragraph_index >= len(paragraphs):
            return {
                "status": "error", 
                "message": f"Paragraph at index {paragraph_index} not found. Document has {len(paragraphs)} paragraphs."
            }

        target_p = paragraphs[paragraph_index]
        
        links.add_hyperlink(
            paragraph=target_p,
            text=text,
            url=url if url else None,
            bookmark=bookmark if bookmark else None
        )
        
        doc.save(file_path)
        
        mode = "external URL" if url else f"internal bookmark '{bookmark}'"
        return {
            "status": "success",
            "message": f"Hyperlink pointing to {mode} added to paragraph {paragraph_index}",
            "file_path": str(file_path)
        }
    except Exception as e:
        return {"status": "error", "message": f"Failed to add hyperlink: {e!s}"}


@validate_docx_write("file_path")
async def add_bookmark(file_path: str, paragraph_index: int, name: str) -> dict[str, Any]:
    """Add a bookmark to a specific paragraph to enable internal linking.

    Args:
        file_path: Path to the .docx file
        paragraph_index: The index of the paragraph to attach the bookmark to (0-based)
        name: The name of the bookmark (should not contain spaces, they will be replaced by underscores)

    Returns:
        A dictionary containing the status and a message or error details
    """
    try:
        doc = Document(file_path)
        
        paragraphs = doc.paragraphs
        if paragraph_index < 0 or paragraph_index >= len(paragraphs):
            return {
                "status": "error", 
                "message": f"Paragraph at index {paragraph_index} not found. Document has {len(paragraphs)} paragraphs."
            }
             
        target_p = paragraphs[paragraph_index]
        links.add_bookmark(target_p, name)
        
        doc.save(file_path)
        
        return {
            "status": "success",
            "message": f"Bookmark '{name}' added to paragraph {paragraph_index}",
            "file_path": str(file_path)
        }
    except Exception as e:
        return {"status": "error", "message": f"Failed to add bookmark: {e!s}"}
