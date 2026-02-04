"""Document creation and manipulation tools for Word Document Server.

This module provides functions for creating, reading, and manipulating
Word documents with security checks for file system access.
"""

import os
from typing import Any

from docx import Document
from docx.document import Document as DocumentType

from mcp_word.core.styles import ensure_heading_style, ensure_table_style
from mcp_word.core.tables import copy_table
from mcp_word.exception import (
    DocumentProcessingError,
    ExceptionTool,
    FileOperationError,
)
from mcp_word.utils.document_utils import (
    create_document_copy,
    ensure_docx_extension,
    extract_document_text,
    get_document_properties,
    get_document_structure,
)
from mcp_word.validation.document_validators import (
    check_file_writeable,
    validate_docx_file,
)


def _get_allowed_directories() -> list[str]:
    """Get the list of allowed directories from environment variables.

    Returns:
        List of absolute paths to directories where documents can be created/accessed.
        Defaults to ['./documents'] if MCP_ALLOWED_DIRECTORIES is not set.
    """

    allowed_dirs_str = os.environ.get("MCP_ALLOWED_DIRECTORIES", "./documents")

    allowed_dirs = [dir.strip() for dir in allowed_dirs_str.split(",")]

    return [os.path.abspath(dir) for dir in allowed_dirs]


def _is_path_in_allowed_directories(file_path: str) -> tuple[bool, str | None]:
    """Check if the given file path is within allowed directories.

    Args:
        file_path: The file path to validate.

    Returns:
        Tuple of (is_allowed, error_message) where is_allowed is a boolean
        indicating if the path is allowed, and error_message provides details
        if the path is not allowed.
    """
    allowed_dirs = _get_allowed_directories()
    abs_path = os.path.abspath(file_path)

    for allowed_dir in allowed_dirs:
        try:
            if os.path.commonpath([allowed_dir, abs_path]) == allowed_dir:
                return True, None
        except ValueError:
            continue

    error_msg = (
        f"Path '{file_path}' is not in allowed directories: {', '.join(allowed_dirs)}"
    )
    return False, error_msg


async def create_document(
    filename: str, title: str | None = None, author: str | None = None
) -> dict[str, Any]:
    """Create a new Word document with optional metadata.

    Args:
        filename: Name of the document to create (with or without .docx extension).
        title: Optional title for the document metadata.
        author: Optional author for the document metadata.

    Returns:
        str: Success or error message indicating the result of the operation.
    """

    is_allowed, error_message = _is_path_in_allowed_directories(filename)
    if not is_allowed:
        return {
            "status": "error",
            "message": f"Cannot create document: {error_message}",
        }

    directory = os.path.dirname(filename)
    if directory and not os.path.exists(directory):
        try:
            os.makedirs(directory, exist_ok=True)
        except OSError as e:
            return ExceptionTool.handle_error(
                FileOperationError(f"Cannot create directory '{directory}': {str(e)}"),
                filename=filename,
                operation="create directory",
            )

    try:
        doc: DocumentType = Document()

        if title:
            doc.core_properties.title = title
        if author:
            doc.core_properties.author = author

        ensure_heading_style(doc)
        ensure_table_style(doc)

        doc.save(filename)

        return {
            "status": "success",
            "message": f"Document {filename} created successfully",
        }
    except Exception as e:
        return ExceptionTool.handle_error(
            DocumentProcessingError(f"Failed to create document: {str(e)}"),
            filename=filename,
            operation="create document",
        )


@validate_docx_file("filename")
async def get_document_info(
    filename: str, response_format: str = "markdown"
) -> dict[str, Any]:
    """Get information about a Word document.

    Args:
        filename: Path to the Word document
        response_format: Format of response - 'markdown' for human-readable or 'json' for structured data
    """
    try:
        properties: dict[str, Any] = get_document_properties(filename)
        return properties
    except Exception as e:
        return ExceptionTool.handle_error(
            DocumentProcessingError(f"Failed to get document info: {str(e)}"),
            filename=filename,
            operation="get document info",
        )


@validate_docx_file("filename")
async def get_document_text(
    filename: str, response_format: str = "markdown"
) -> dict[str, Any]:
    """Extract all text from a Word document.

    Args:
        filename: Path to the Word document
        response_format: Format of response - 'markdown' for human-readable or 'json' for structured data
    """
    try:
        result: dict[str, Any] = extract_document_text(filename)
        return result
    except Exception as e:
        return ExceptionTool.handle_error(
            DocumentProcessingError(f"Failed to get document text: {str(e)}"),
            filename=filename,
            operation="get document text",
        )


@validate_docx_file("filename")
async def get_document_outline(
    filename: str, response_format: str = "markdown"
) -> dict[str, Any]:
    """Get the structure of a Word document.

    Args:
        filename: Path to the Word document
        response_format: Format of response - 'markdown' for human-readable or 'json' for structured data
    """
    try:
        structure: dict[str, Any] = get_document_structure(filename)
        return structure
    except Exception as e:
        return ExceptionTool.handle_error(
            DocumentProcessingError(f"Failed to get document outline: {str(e)}"),
            filename=filename,
            operation="get document outline",
        )


async def list_available_documents(
    directory: str | None = None,
    page: int = 1,
    page_size: int = 20,
    response_format: str = "markdown",
) -> dict[str, Any]:
    """List all .docx files in the allowed directories with pagination support.

    Args:
        directory: Optional specific directory to search (defaults to all allowed directories)
        page: Page number for pagination (1-based, default: 1)
        page_size: Number of documents per page (default: 20, max: 100)
        response_format: Format of response - 'markdown' for human-readable or 'json' for structured data

    Returns:
        Dictionary with status, message, pagination info, and list of found documents.
    """
    try:
        # Validate pagination parameters
        if page < 1:
            page = 1
        if page_size < 1:
            page_size = 20
        if page_size > 100:
            page_size = 100

        search_directories = _get_allowed_directories()
        if not search_directories:
            return {
                "status": "error",
                "message": "No accessible directories found in MCP configuration.",
                "allowed_directories": search_directories,
                "documents": [],
                "pagination": {
                    "page": page,
                    "page_size": page_size,
                    "total": 0,
                    "has_more": False,
                    "next_offset": None,
                },
            }

        all_documents = []
        total_found = 0

        for search_dir in search_directories:
            if not os.path.exists(search_dir):
                continue

            docx_files = [
                f
                for f in os.listdir(search_dir)
                if f.lower().endswith(".docx") and not f.startswith("~$")
            ]

            for file in sorted(docx_files):
                file_path = os.path.join(search_dir, file)
                size_kb = os.path.getsize(file_path) / 1024
                all_documents.append(
                    {
                        "name": file,
                        "path": os.path.abspath(file_path),
                        "size_kb": round(size_kb, 2),
                        "source_directory": search_dir,
                    }
                )

            total_found += len(docx_files)

        # Apply pagination
        total_documents = len(all_documents)
        offset = (page - 1) * page_size
        paginated_documents = all_documents[offset : offset + page_size]
        has_more = offset + page_size < total_documents
        next_offset = page + 1 if has_more else None

        message = f"Found {total_found} Word document(s)"
        if len(search_directories) > 1:
            message += f" across {len(search_directories)} directories"
        message += f". Showing page {page} of {((total_documents - 1) // page_size) + 1 if total_documents > 0 else 1}"

        return {
            "status": "success",
            "message": message,
            "directories_searched": search_directories,
            "total": total_found,
            "documents": paginated_documents,
            "pagination": {
                "page": page,
                "page_size": page_size,
                "total": total_documents,
                "has_more": has_more,
                "next_offset": next_offset,
            },
            "response_format": response_format,
        }

    except Exception as e:
        return ExceptionTool.handle_error(
            FileOperationError(f"Failed to list documents: {str(e)}"),
            operation="list documents",
        )


@validate_docx_file("source_filename")
async def copy_document(
    source_filename: str, destination_filename: str | None = None
) -> dict[str, Any]:
    """Create a copy of a Word document.

    Args:
        source_filename: Path to the source document
        destination_filename: Optional path for the copy. If not provided, a default name will be generated.
    """
    source_filename = ensure_docx_extension(source_filename)

    if destination_filename:
        destination_filename = ensure_docx_extension(destination_filename)

    success, message, _ = create_document_copy(source_filename, destination_filename)
    if success:
        return {"status": "success", "message": message}
    return ExceptionTool.handle_error(
        FileOperationError(f"Failed to copy document: {message}"),
        filename=source_filename,
        operation="copy document",
    )


@validate_docx_file("target_filename")
@check_file_writeable("target_filename")
async def merge_documents(
    target_filename: str,
    source_filenames: list[str],
    add_page_breaks: bool = True,
    response_format: str = "markdown",
) -> dict[str, Any]:
    """Merge multiple Word documents into a single document.

    Args:
        target_filename: Path to the target document (will be created or overwritten)
        source_filenames: List of paths to source documents to merge
        add_page_breaks: If True, add page breaks between documents
        response_format: Format of response - 'markdown' for human-readable or 'json' for structured data

    Returns:
        Dictionary with status, message, and details about merged documents
    """
    # Validate all source documents exist
    missing_files = []
    for filename in source_filenames:
        doc_filename = ensure_docx_extension(filename)
        if not os.path.exists(doc_filename):
            missing_files.append(doc_filename)

    if missing_files:
        return {
            "status": "error",
            "message": f"Cannot merge documents. The following source files do not exist: {', '.join(missing_files)}",
        }

    try:
        # Create a new document for the merged result
        target_doc: DocumentType = Document()

        # Process each source document
        for i, filename in enumerate(source_filenames):
            doc_filename = ensure_docx_extension(filename)
            source_doc: DocumentType = Document(doc_filename)

            # Add page break between documents (except before the first one)
            if add_page_breaks and i > 0:
                target_doc.add_page_break()

            # Copy all paragraphs
            for paragraph in source_doc.paragraphs:
                # Create a new paragraph with the same text and style
                new_paragraph = target_doc.add_paragraph(paragraph.text)
                new_paragraph.style = target_doc.styles["Normal"]  # Default style

                # Try to match the style if possible
                try:
                    if paragraph.style and paragraph.style.name in target_doc.styles:
                        new_paragraph.style = target_doc.styles[paragraph.style.name]
                except:  # noqa: E722
                    pass

                # Copy run formatting
                for i, run in enumerate(paragraph.runs):
                    if i < len(new_paragraph.runs):
                        new_run = new_paragraph.runs[i]
                        # Copy basic formatting
                        new_run.bold = run.bold
                        new_run.italic = run.italic
                        new_run.underline = run.underline
                        # Font size if specified
                        if run.font.size:
                            new_run.font.size = run.font.size

            # Copy all tables
            for table in source_doc.tables:
                copy_table(table, target_doc)

        # Save the merged document
        target_doc.save(target_filename)
        return {
            "status": "success",
            "message": f"Successfully merged {len(source_filenames)} documents into {target_filename}",
        }
    except Exception as e:
        return ExceptionTool.handle_error(
            DocumentProcessingError(f"Failed to merge documents: {str(e)}"),
            filename=target_filename,
            operation="merge documents",
        )
