"""Formatting tools for Word Document Server.

These tools handle formatting operations for Word documents,
including text formatting, table formatting, and custom styles.
"""

from typing import Any
from docx.enum.style import WD_STYLE_TYPE

from mcp_word.core import (
    core_format_text,
    create_style,
    apply_table_style,
)
from mcp_word.core.document_context import document_context
from mcp_word.exception import (
    DocumentProcessingError,
    ExceptionTool,
    StyleError,
)
from mcp_word.models.response_models import (
    TextFormatResult,
    StyleResult,
    TableFormatResult,
)
from mcp_word.validation.document_validators import validate_docx_write


@validate_docx_write("filename")
async def format_text(
    filename: str,
    paragraph_index: int,
    start_pos: int,
    end_pos: int,
    bold: bool | None = None,
    italic: bool | None = None,
    underline: bool | None = None,
    color: str | None = None,
    font_size: int | None = None,
    font_name: str | None = None,
) -> dict[str, Any]:
    """Format a specific range of text within a paragraph."""
    try:
        # Parameter validation
        p_index = int(paragraph_index)
        s_pos = int(start_pos)
        e_pos = int(end_pos)
        f_size = int(font_size) if font_size is not None else None

        with document_context(filename, mode="write") as doc:
            result = core_format_text(
                doc,
                p_index,
                s_pos,
                e_pos,
                bold=bold,
                italic=italic,
                underline=underline,
                color=color,
                font_size=f_size,
                font_name=font_name,
            )

        return TextFormatResult(
            status="success",
            filename=filename,
            paragraph_index=p_index,
            text_formatted=result["text_formatted"],
            formatting_applied=result["formatting_applied"],
            message=f"Text formatted successfully in paragraph {p_index}.",
        ).model_dump()

    except (ValueError, TypeError, IndexError, DocumentProcessingError, OSError) as error:
        return ExceptionTool.handle_error(
            error if isinstance(error, DocumentProcessingError)
            else DocumentProcessingError(f"Failed to format text: {str(error)}"),
            filename=filename,
            operation="format text",
        )


@validate_docx_write("filename")
async def create_custom_style(
    filename: str,
    style_name: str,
    bold: bool | None = None,
    italic: bool | None = None,
    font_size: int | None = None,
    font_name: str | None = None,
    color: str | None = None,
    base_style: str | None = None,
) -> dict[str, Any]:
    """Create a custom style in the document."""
    try:
        # Build font properties dictionary
        font_properties: dict[str, Any] = {}
        if bold is not None:
            font_properties["bold"] = bold
        if italic is not None:
            font_properties["italic"] = italic
        if font_size is not None:
            font_properties["size"] = font_size
        if font_name is not None:
            font_properties["name"] = font_name
        if color is not None:
            font_properties["color"] = color

        with document_context(filename, mode="write") as doc:
            create_style(
                doc,
                style_name,
                WD_STYLE_TYPE.PARAGRAPH,
                base_style=base_style,
                font_props=font_properties,
            )

        return StyleResult(
            status="success",
            filename=filename,
            style_name=style_name,
            style_type="PARAGRAPH",
            font_properties=font_properties,
            message=f"Style '{style_name}' created successfully.",
        ).model_dump()

    except (OSError, ValueError, KeyError, StyleError) as error:
        return ExceptionTool.handle_error(
            error if isinstance(error, StyleError)
            else StyleError(f"Failed to create style: {str(error)}"),
            filename=filename,
            operation="create custom style",
        )


@validate_docx_write("filename")
async def format_table(
    filename: str,
    table_index: int,
    has_header_row: bool | None = None,
    border_style: str | None = None,
    shading: list[list[str]] | None = None,
) -> dict[str, Any]:
    """Format a table with borders, shading, and structure."""
    try:
        t_index = int(table_index)
        
        with document_context(filename, mode="write") as doc:
            if t_index < 0 or t_index >= len(doc.tables):
                raise IndexError(f"Invalid table index {t_index}. Document has {len(doc.tables)} tables.")
            
            table = doc.tables[t_index]
            
            # Apply formatting
            table_shading: list[list[str | None]] | None = None
            if shading is not None:
                table_shading = [list(row) for row in shading]

            success = apply_table_style(
                table, has_header_row or False, border_style, table_shading
            )

            if not success:
                raise DocumentProcessingError(f"Failed to apply table style to table at index {t_index}")

        return TableFormatResult(
            status="success",
            filename=filename,
            table_index=t_index,
            header_row_applied=has_header_row or False,
            border_style=border_style,
            shading_applied=shading is not None,
            message=f"Table at index {t_index} formatted successfully.",
        ).model_dump()

    except (OSError, ValueError, KeyError, IndexError, DocumentProcessingError) as error:
        return ExceptionTool.handle_error(
            error if isinstance(error, DocumentProcessingError)
            else DocumentProcessingError(f"Failed to format table: {str(error)}"),
            filename=filename,
            operation="format table",
        )
