"""FastMCP tools for document properties and page layout."""

from typing import Any

from docx import Document

from mcp_word.core import properties
from mcp_word.validation.document_validators import validate_docx_read, validate_docx_write


@validate_docx_read("file_path")
async def get_core_properties(file_path: str) -> dict[str, Any]:
    """Get the core metadata properties of the document (author, title, etc.).

    Args:
        file_path: Path to the .docx file

    Returns:
        A dictionary containing the document properties
    """
    try:
        doc = Document(file_path)
        
        props = properties.get_core_properties(doc)
        
        return {
            "status": "success",
            "properties": props,
        }
    except Exception as e:
        return {"status": "error", "message": f"Failed to get properties: {e!s}"}


@validate_docx_write("file_path")
async def set_core_properties(
    file_path: str,
    author: str | None = None,
    title: str | None = None,
    subject: str | None = None,
    keywords: str | None = None,
    comments: str | None = None,
    category: str | None = None
) -> dict[str, Any]:
    """Set the core metadata properties of the document.

    Args:
        file_path: Path to the .docx file
        author: Optional author name
        title: Optional document title
        subject: Optional document subject
        keywords: Optional document keywords/tags
        comments: Optional document comments
        category: Optional document category

    Returns:
        A dictionary containing the status and a message or error details
    """
    try:
        doc = Document(file_path)
        
        # Build kwargs dictionary excluding None values
        kwargs = {}
        if author is not None: kwargs["author"] = author
        if title is not None: kwargs["title"] = title
        if subject is not None: kwargs["subject"] = subject
        if keywords is not None: kwargs["keywords"] = keywords
        if comments is not None: kwargs["comments"] = comments
        if category is not None: kwargs["category"] = category
        
        if not kwargs:
            return {"status": "error", "message": "No properties provided to update."}
            
        properties.set_core_properties(doc, **kwargs)
        doc.save(file_path)
        
        return {
            "status": "success",
            "message": f"Successfully updated properties: {list(kwargs.keys())}",
            "file_path": str(file_path)
        }
    except Exception as e:
        return {"status": "error", "message": f"Failed to set properties: {e!s}"}


@validate_docx_write("file_path")
async def set_page_layout(file_path: str, section_index: int = 0, orientation: str = "portrait") -> dict[str, Any]:
    """Set the page layout orientation (portrait or landscape) for a specific section.

    Args:
        file_path: Path to the .docx file
        section_index: The index of the section (0-based). Default is 0.
        orientation: Must be either "portrait" or "landscape".

    Returns:
        A dictionary containing the status and a message or error details
    """
    try:
        doc = Document(file_path)
        
        properties.set_page_layout(doc, section_idx=section_index, orientation=orientation)
        doc.save(file_path)
        
        return {
            "status": "success",
            "message": f"Successfully set section {section_index} layout to {orientation}",
            "file_path": str(file_path)
        }
    except ValueError as e:
        return {"status": "error", "message": str(e)}
    except IndexError as e:
        return {"status": "error", "message": str(e)}
    except Exception as e:
        return {"status": "error", "message": f"Failed to set page layout: {e!s}"}
