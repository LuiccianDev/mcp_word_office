"""Formatting-related core operations for Word Document Server.

This module provides functions for formatting text, paragraphs, and other
elements within Word documents.
"""

from typing import Any
from docx.document import Document as DocumentType
from docx.shared import Pt, RGBColor
from docx.text.paragraph import Paragraph

def core_format_text(
    doc: DocumentType,
    paragraph_index: int,
    start_pos: int,
    end_pos: int,
    bold: bool | None = None,
    italic: bool | None = None,
    underline: bool | None = None,
    color: str | None = None,
    font_size: int | None = None,
    font_name: str | None = None,
) -> dict[str, Any]:
    """Format a specific range of text within a paragraph.

    Args:
        doc: The document object.
        paragraph_index: Index of the paragraph (0-based).
        start_pos: Start position within the paragraph text.
        end_pos: End position within the paragraph text.
        bold: Set text bold.
        italic: Set text italic.
        underline: Set text underlined.
        color: Text color (name or hex).
        font_size: Font size in points.
        font_name: Font name/family.

    Returns:
        dict: Details of the formatting applied.
    """
    if paragraph_index < 0 or paragraph_index >= len(doc.paragraphs):
        raise IndexError(f"Paragraph index {paragraph_index} is out of range.")

    paragraph = doc.paragraphs[paragraph_index]
    text = paragraph.text

    if start_pos < 0 or end_pos > len(text) or start_pos >= end_pos:
        raise ValueError(f"Invalid text positions: {start_pos}-{end_pos} for text of length {len(text)}")

    target_text = text[start_pos:end_pos]

    # Clear existing runs and create three runs: before, target, after
    for run in paragraph.runs:
        run.clear()

    # Add text before target
    if start_pos > 0:
        paragraph.add_run(text[:start_pos])

    # Add target text with formatting
    run_target = paragraph.add_run(target_text)
    
    formatting_applied = {}
    
    if bold is not None:
        run_target.bold = bold
        formatting_applied["bold"] = bold
    if italic is not None:
        run_target.italic = italic
        formatting_applied["italic"] = italic
    if underline is not None:
        run_target.underline = underline
        formatting_applied["underline"] = underline
        
    if color:
        color_map = {
            "red": RGBColor(255, 0, 0),
            "blue": RGBColor(0, 0, 255),
            "green": RGBColor(0, 128, 0),
            "yellow": RGBColor(255, 255, 0),
            "black": RGBColor(0, 0, 0),
            "gray": RGBColor(128, 128, 128),
            "white": RGBColor(255, 255, 255),
            "purple": RGBColor(128, 0, 128),
            "orange": RGBColor(255, 165, 0),
        }

        try:
            if color.lower() in color_map:
                run_target.font.color.rgb = color_map[color.lower()]
                formatting_applied["color"] = color.lower()
            else:
                run_target.font.color.rgb = RGBColor.from_string(color.replace("#", ""))
                formatting_applied["color"] = color
        except (ValueError, AttributeError):
            run_target.font.color.rgb = RGBColor(0, 0, 0)
            formatting_applied["color"] = "default (black)"
            
    if font_size:
        run_target.font.size = Pt(font_size)
        formatting_applied["font_size"] = font_size
    if font_name:
        run_target.font.name = font_name
        formatting_applied["font_name"] = font_name

    # Add text after target
    if end_pos < len(text):
        paragraph.add_run(text[end_pos:])

    return {
        "paragraph_index": paragraph_index,
        "text_formatted": target_text,
        "formatting_applied": formatting_applied
    }
