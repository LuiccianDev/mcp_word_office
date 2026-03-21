"""
Footnote and endnote functionality for Word Document Server.
"""

import string
from typing import Any

from docx.document import Document
from docx.text.paragraph import Paragraph


def core_add_footnote(doc: Document, paragraph_index: int, text: str) -> dict[str, Any]:
    """Add a footnote to a specific paragraph.

    Args:
        doc: Document object.
        paragraph_index: Index of the paragraph.
        text: Text content of the footnote.

    Returns:
        dict: Details of the added footnote.
    """
    if paragraph_index < 0 or paragraph_index >= len(doc.paragraphs):
        raise IndexError(f"Paragraph index {paragraph_index} out of range.")

    paragraph = doc.paragraphs[paragraph_index]
    
    # Superscript reference
    run = paragraph.add_run()
    run.text = "¹"
    run.font.superscript = True

    # Add footnote section at the end if it doesn't exist
    found = False
    for p in doc.paragraphs:
        if p.text.startswith("Footnotes:"):
            found = True
            break
    if not found:
        doc.add_paragraph("\n")
        doc.add_paragraph().add_run("Footnotes:").bold = True

    doc.add_paragraph(f"¹ {text}")

    return {
        "footnote_id": 1, # Placeholder as docx doesn't support real IDs easily
        "footnote_text": text,
        "paragraph_index": paragraph_index
    }


def core_add_endnote(doc: Document, paragraph_index: int, text: str) -> dict[str, Any]:
    """Add an endnote to a specific paragraph.

    Args:
        doc: Document object.
        paragraph_index: Index of the paragraph.
        text: Text content of the endnote.

    Returns:
        dict: Details of the added endnote.
    """
    if paragraph_index < 0 or paragraph_index >= len(doc.paragraphs):
        raise IndexError(f"Paragraph index {paragraph_index} out of range.")

    paragraph = doc.paragraphs[paragraph_index]
    
    run = paragraph.add_run()
    run.text = "*"
    run.font.superscript = True

    # Endnotes section
    found = False
    for para in doc.paragraphs:
        if para.text == "Endnotes:":
            found = True
            break

    if not found:
        doc.add_page_break()
        doc.add_heading("Endnotes:", level=1)

    doc.add_paragraph(f"* {text}")

    return {
        "endnote_id": 1,
        "endnote_text": text,
        "paragraph_index": paragraph_index
    }


def core_convert_footnotes_to_endnotes(doc: Document) -> int:
    """
    Convert all footnotes to endnotes in a document.

    Args:
        doc: Document object

    Returns:
        Number of footnotes converted
    """
    # This is a complex operation not fully supported by python-docx
    # Implementing a simplified version

    # Collect all footnotes
    footnotes = []
    for para in doc.paragraphs:
        # This is a simplified implementation
        for run in para.runs:
            if run.font.superscript and run.text.isdigit():
                # This might be a footnote reference
                footnotes.append((para, run.text))

    # Add endnotes section
    if footnotes:
        doc.add_page_break()
        doc.add_heading("Endnotes:", level=1)

        # Add each footnote as an endnote
        for idx, (_para, footnote_num) in enumerate(footnotes):
            doc.add_paragraph(f"{idx + 1}. Converted from footnote {footnote_num}")

    return len(footnotes)


def find_footnote_references(doc: Document) -> list[tuple[int, int, str]]:
    """
    Find all footnote references in a document.

    Args:
        doc: Document object

    Returns:
        List of tuples (paragraph_index, run_index, text) for each footnote reference
    """
    footnote_references: list[tuple[int, int, str]] = []

    for para_idx, para in enumerate(doc.paragraphs):
        for run_idx, run in enumerate(para.runs):
            if run.font.superscript and (run.text.isdigit() or run.text in "¹²³⁴⁵⁶⁷⁸⁹"):
                footnote_references.append((para_idx, run_idx, run.text))

    return footnote_references


def _fill_or_extend_format_symbols(base_list: list[str], count: int) -> list[str]:
    """
    Returns a list of length `count` based on `base_list`.
    If `count` exceeds `base_list`, fills the rest with numeric strings.
    """
    if count <= len(base_list):
        return base_list[:count]
    return base_list + [str(i) for i in range(1, count - len(base_list) + 1)]


def get_format_symbols(numbering_format: str, count: int) -> list[str]:
    """
    Get a list of formatting symbols based on the specified numbering format.

    Args:
        numbering_format: Format for footnote/endnote numbers (e.g., "1, 2, 3", "i, ii, iii", "a, b, c")
        count: Number of symbols needed

    Returns:
        List of formatting symbols
    """
    roman_numerals = [
        "i",
        "ii",
        "iii",
        "iv",
        "v",
        "vi",
        "vii",
        "viii",
        "ix",
        "x",
        "xi",
        "xii",
        "xiii",
        "xiv",
        "xv",
        "xvi",
        "xvii",
        "xviii",
        "xix",
        "xx",
    ]
    alphabet = list(string.ascii_lowercase)
    symbols = ["*", "†", "‡", "§", "¶", "||", "**", "††", "‡‡", "§§"]

    if numbering_format == "i, ii, iii":
        return _fill_or_extend_format_symbols(roman_numerals, count)
    elif numbering_format == "a, b, c":
        return _fill_or_extend_format_symbols(alphabet, count)
    elif numbering_format == "*, †, ‡":
        return _fill_or_extend_format_symbols(symbols, count)
    else:  # Default to numeric strings
        return [str(i) for i in range(1, count + 1)]


def core_customize_footnote_formatting(
    doc: Document,
    footnote_refs: list[tuple[int, int, str]],
    format_symbols: list[str],
    start_number: int,
    style: str | None = None,
) -> int:
    """
    Apply custom formatting to footnote references and text.

    Args:
        doc: Document object
        footnote_refs: List of footnote references from find_footnote_references()
        format_symbols: List of formatting symbols to use
        start_number: Number to start footnote numbering from
        style: Optional style to apply to footnote text

    Returns:
        Number of footnotes formatted
    """
    # Update footnote references with new format
    for i, (para_idx, run_idx, _) in enumerate(footnote_refs):
        try:
            idx = i + start_number - 1
            if idx < len(format_symbols):
                symbol = format_symbols[idx]
            else:
                symbol = str(idx + 1)  # Fall back to numbers if we run out of symbols

            paragraph = doc.paragraphs[para_idx]
            paragraph.runs[run_idx].text = symbol
        except IndexError:
            # Skip if we can't locate the reference
            pass

    # Find footnote section and update
    for para_idx, para in enumerate(doc.paragraphs):
        if para.text.startswith("Footnotes:") or para.text == "Footnotes":
            # Update footnotes with new symbols
            for i in range(len(footnote_refs)):
                try:
                    footnote_para_idx = para_idx + i + 1
                    if footnote_para_idx < len(doc.paragraphs):
                        para = doc.paragraphs[footnote_para_idx]

                        # Extract and preserve footnote text
                        footnote_text = para.text
                        if " " in footnote_text and len(footnote_text) > 2:
                            # Remove the old footnote number/symbol
                            footnote_text = footnote_text.split(" ", 1)[1]

                        # Add new format
                        idx = i + start_number - 1
                        if idx < len(format_symbols):
                            symbol = format_symbols[idx]
                        else:
                            symbol = str(idx + 1)

                        # Apply new formatting
                        para.text = f"{symbol} {footnote_text}"

                        # Apply style
                        if style:
                            para.style = style
                except IndexError:
                    pass

            break

    return len(footnote_refs)
