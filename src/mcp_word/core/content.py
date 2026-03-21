"""Content-related core operations for Word Document Server.

This module provides functions to manipulate various types of content
in Word documents, including headings, paragraphs, and images.
"""

import os
from typing import Any

from docx.document import Document as DocumentType
from docx.shared import Inches, Pt
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.text import WD_ALIGN_PARAGRAPH

from mcp_word.core.styles import ensure_heading_style


def core_add_heading(doc: DocumentType, text: str, level: int = 1) -> dict[str, Any]:
    """Add a heading to a Word document.

    Args:
        doc: The document object.
        text: The text content of the heading.
        level: The heading level (1-9).

    Returns:
        dict: Details of the added heading.
    """
    try:
        ensure_heading_style(doc)
        doc.add_heading(text, level=level)
        return {"heading_text": text, "heading_level": level}
    except (KeyError, AttributeError):
        # Fallback for missing styles
        paragraph = doc.add_paragraph(text)
        paragraph.style = doc.styles["Normal"]
        run = paragraph.runs[0]
        run.bold = True
        font_sizes = {1: 16, 2: 14}
        run.font.size = Pt(font_sizes.get(level, 12))
        return {"heading_text": text, "heading_level": level, "fallback_applied": True}


def core_add_paragraph(doc: DocumentType, text: str, style: str | None = None) -> dict[str, Any]:
    """Add a paragraph to a Word document.

    Args:
        doc: The document object.
        text: The text content.
        style: Optional style name.

    Returns:
        dict: Details of the added paragraph.
    """
    paragraph = doc.add_paragraph(text)
    style_applied = "Normal"

    if style:
        try:
            paragraph.style = style
            style_applied = style
        except KeyError:
            paragraph.style = doc.styles["Normal"]
            return {"paragraph_text": text, "style_applied": "Normal", "style_requested_missing": style}

    return {"paragraph_text": text, "style_applied": style_applied}


def core_add_picture(doc: DocumentType, image_path: str, width: float | None = None) -> dict[str, Any]:
    """Add an image to a Word document.

    Args:
        doc: The document object.
        image_path: Absolute path to the image file.
        width: Optional width in inches.

    Returns:
        dict: Details of the added picture.
    """
    if width is not None:
        doc.add_picture(image_path, width=Inches(width))
    else:
        doc.add_picture(image_path)
        
    return {
        "image_name": os.path.basename(image_path),
        "image_path": image_path,
        "width_inches": width
    }


def core_add_page_break(doc: DocumentType) -> None:
    """Add a page break to the document."""
    doc.add_page_break()


def core_delete_paragraph(doc: DocumentType, paragraph_index: int) -> None:
    """Delete a paragraph from a document."""
    total_paragraphs = len(doc.paragraphs)
    if paragraph_index < 0 or paragraph_index >= total_paragraphs:
        raise IndexError(f"Paragraph index {paragraph_index} is out of range (0-{total_paragraphs - 1})")

    paragraph = doc.paragraphs[paragraph_index]
    paragraph_element = paragraph._p
    paragraph_element.getparent().remove(paragraph_element)


def core_add_table_of_contents(
    doc: DocumentType, title: str = "Table of Contents", max_level: int = 3
) -> dict[str, Any]:
    """Add a Table of Contents to the document.

    Note: The TOC needs to be updated manually in Word or via a macro.

    Args:
        doc: The document object.
        title: Title of the TOC.
        max_level: Maximum heading level to include.

    Returns:
        dict: Details of the added TOC.
    """
    paragraph = doc.add_paragraph()
    run = paragraph.add_run(title)
    run.bold = True
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    paragraph = doc.add_paragraph()
    run = paragraph.add_run()

    # Create the TOC fields
    fldChar = OxmlElement("w:fldChar")  # noqa: N806
    fldChar.set(qn("w:fldCharType"), "begin")

    instrText = OxmlElement("w:instrText")  # noqa: N806
    instrText.set(qn("xml:space"), "preserve")
    instrText.text = f'TOC \\o "1-{max_level}" \\h \\z \\u'

    fldChar2 = OxmlElement("w:fldChar")  # noqa: N806
    fldChar2.set(qn("w:fldCharType"), "separate")

    fldChar3 = OxmlElement("w:fldChar")  # noqa: N806
    fldChar3.set(qn("w:fldCharType"), "end")

    run._r.append(fldChar)
    run._r.append(instrText)
    run._r.append(fldChar2)
    run._r.append(fldChar3)

    doc.add_page_break()

    return {
        "title": title,
        "max_level": max_level,
        "message": "Table of Contents field added. The document may need manual update in Word to display properly.",
    }


def core_find_and_replace_text(doc: DocumentType, old_text: str, new_text: str) -> int:
    """Find and replace text throughout the document.

    Args:
        doc: Document object to search and modify.
        old_text: Text to find.
        new_text: Text to replace with.

    Returns:
        Number of replacements made.
    """
    if not old_text:
        raise ValueError("Search text cannot be empty")

    replacement_count: int = 0

    # Search and replace in paragraphs
    for paragraph in doc.paragraphs:
        if old_text in paragraph.text:
            for run in paragraph.runs:
                if old_text in run.text:
                    run.text = run.text.replace(old_text, new_text)
                    replacement_count += 1

    # Search and replace in tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    if old_text in paragraph.text:
                        for run in paragraph.runs:
                            if old_text in run.text:
                                run.text = run.text.replace(old_text, new_text)
                                replacement_count += 1

    return replacement_count
