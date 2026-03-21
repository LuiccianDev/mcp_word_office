"""
Link and bookmark operations for Word Document Server.

This module provides functionality for adding hyperlinks (internal and external)
and bookmarks to document paragraphs.
"""

from docx.document import Document
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.oxml.ns import qn
from docx.oxml.parser import OxmlElement
from docx.text.paragraph import Paragraph
from docx.text.run import Run


def core_add_bookmark(doc: Document, paragraph_index: int, name: str) -> None:
    """Add a bookmark to the end of a paragraph.

    Args:
        doc: The Word document object.
        paragraph_index: Index of the paragraph.
        name: The name of the bookmark.
    """
    if paragraph_index < 0 or paragraph_index >= len(doc.paragraphs):
        raise IndexError(f"Paragraph index {paragraph_index} out of range.")
    
    paragraph = doc.paragraphs[paragraph_index]
    
    # Clean the name to be a valid xml id (no spaces)
    name = name.replace(" ", "_")
    
    # Create the xml elements
    bookmark_start = OxmlElement("w:bookmarkStart")
    bookmark_start.set(qn("w:id"), "0")
    bookmark_start.set(qn("w:name"), name)

    bookmark_end = OxmlElement("w:bookmarkEnd")
    bookmark_end.set(qn("w:id"), "0")

    # Append to the paragraph element
    paragraph._p.append(bookmark_start)
    paragraph._p.append(bookmark_end)


def core_add_hyperlink(
    doc: Document, paragraph_index: int, text: str, url: str | None = None, bookmark: str | None = None
) -> None:
    """Add a hyperlink into a paragraph.

    Args:
        doc: The Word document object.
        paragraph_index: Index of the paragraph.
        text: The clickable text of the hyperlink.
        url: The external URL.
        bookmark: The internal bookmark name to link to.
    """
    if paragraph_index < 0 or paragraph_index >= len(doc.paragraphs):
        raise IndexError(f"Paragraph index {paragraph_index} out of range.")
        
    paragraph = doc.paragraphs[paragraph_index]

    if bool(url) == bool(bookmark):
        raise ValueError("Must provide exactly one of 'url' or 'bookmark'")

    # Create the w:hyperlink tag
    hyperlink = OxmlElement("w:hyperlink")

    if url:
        # Create a relationship for external link
        part = paragraph.part
        r_id = part.relate_to(url, RT.HYPERLINK, is_external=True)
        hyperlink.set(qn("r:id"), r_id)
    elif bookmark:
        # Link to internal bookmark
        bookmark_cleaned = bookmark.replace(" ", "_")
        hyperlink.set(qn("w:anchor"), bookmark_cleaned)

    # Create a new run with the text
    new_run = Run(OxmlElement("w:r"), paragraph)
    new_run.text = text
    
    # Optional: styling
    new_run.font.color.rgb = None
    new_run.font.color.theme_color = 10 # Hyperlink theme color typically
    new_run.font.underline = True

    # Append the run xml to the hyperlink xml
    hyperlink.append(new_run._r)
    
    # Append the hyperlink to the paragraph xml
    paragraph._p.append(hyperlink)
