"""
Header and footer operations for Word Document Server.

This module provides functionality for working with headers and footers
in document sections.
"""

from docx.document import Document
from docx.enum.section import WD_HEADER_FOOTER


def set_section_header(
    doc: Document, text: str, section_idx: int = 0, is_primary: bool = True
) -> None:
    """Set the header text for a specific section.

    Args:
        doc: The Word document object.
        text: The text to add to the header.
        section_idx: The index of the section to modify (default is 0 for the first section).
        is_primary: Whether to modify the primary header. Currently only primary is supported via simple API.

    Raises:
        IndexError: If the section index is out of bounds.
    """
    sections = doc.sections
    if section_idx < 0 or section_idx >= len(sections):
        raise IndexError(f"Section index {section_idx} is out of bounds. Document has {len(sections)} sections.")

    section = sections[section_idx]
    
    # In python-docx, section.header is the primary header.
    # To modify first page or even page headers, different properties exist.
    header = section.header
    
    # Prevent inherited headers if we're explicitly setting one
    header.is_linked_to_previous = False

    # Clear existing content and set new text
    if len(header.paragraphs) > 0:
        header.paragraphs[0].text = text
    else:
        header.add_paragraph(text)


def set_section_footer(
    doc: Document, text: str, section_idx: int = 0, is_primary: bool = True
) -> None:
    """Set the footer text for a specific section.

    Args:
        doc: The Word document object.
        text: The text to add to the footer.
        section_idx: The index of the section to modify.
        is_primary: Whether to modify the primary footer.

    Raises:
        IndexError: If the section index is out of bounds.
    """
    sections = doc.sections
    if section_idx < 0 or section_idx >= len(sections):
        raise IndexError(f"Section index {section_idx} is out of bounds. Document has {len(sections)} sections.")

    section = sections[section_idx]
    
    footer = section.footer
    footer.is_linked_to_previous = False

    # Clear existing content and set new text
    if len(footer.paragraphs) > 0:
        footer.paragraphs[0].text = text
    else:
        footer.add_paragraph(text)
