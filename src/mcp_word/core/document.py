"""
Core document operations for Word Document Server.

This module provides functionality for document lifecycle management,
including creation, copying, merging, and conversion.
"""

import os
import platform
import shutil
import subprocess
import re
from re import Pattern
from typing import Any

from docx import Document
from docx.document import Document as DocumentType
from docx.text.paragraph import Paragraph

from mcp_word.core.styles import (
    DEFAULT_SETTINGS,
    StyleSettings,
    ensure_heading_style,
    ensure_table_style,
)
from mcp_word.core.tables import copy_table
from mcp_word.utils.document_utils import (
    create_document_copy,
    ensure_docx_extension,
)


def core_create_document(
    filename: str,
    title: str | None = None,
    author: str | None = None,
    style_settings: dict[str, Any] | StyleSettings | None = None,
) -> str:
    """Create a new Word document with optional metadata.

    Args:
        filename: Path to the document to create.
        title: Optional title metadata.
        author: Optional author metadata.

    Returns:
        The path to the created document.
    """
    doc: DocumentType = Document()

    if title:
        doc.core_properties.title = title
    if author:
        doc.core_properties.author = author

    # Initialize style settings
    if isinstance(style_settings, dict):
        settings = StyleSettings(**style_settings)
    elif isinstance(style_settings, StyleSettings):
        settings = style_settings
    else:
        settings = DEFAULT_SETTINGS

    ensure_heading_style(doc, settings=settings)
    ensure_table_style(doc)

    doc.save(filename)
    return filename


def core_copy_document(
    source_filename: str, destination_filename: str | None = None
) -> tuple[bool, str, str]:
    """Create a copy of a Word document.

    Args:
        source_filename: Path to the source document.
        destination_filename: Optional path for the copy.

    Returns:
        Tuple of (success, message, destination_path).
    """
    source_filename = ensure_docx_extension(source_filename)
    if destination_filename:
        destination_filename = ensure_docx_extension(destination_filename)

    return create_document_copy(source_filename, destination_filename)


def core_merge_documents(
    target_filename: str,
    source_filenames: list[str],
    add_page_breaks: bool = True,
) -> str:
    """Merge multiple Word documents into a single document.

    Args:
        target_filename: Path to the target document.
        source_filenames: List of source document paths.
        add_page_breaks: Whether to add page breaks between documents.

    Returns:
        The path to the merged document.
    """
    target_doc: DocumentType = Document()

    for i, filename in enumerate(source_filenames):
        doc_filename = ensure_docx_extension(filename)
        source_doc: DocumentType = Document(doc_filename)

        if add_page_breaks and i > 0:
            target_doc.add_page_break()

        # Copy paragraphs
        for paragraph in source_doc.paragraphs:
            new_paragraph = target_doc.add_paragraph(paragraph.text)
            
            # Simple style matching
            try:
                if paragraph.style and paragraph.style.name in target_doc.styles:
                    new_paragraph.style = paragraph.style.name
            except (KeyError, AttributeError):
                pass

            # Copy run formatting
            for run_idx, run in enumerate(paragraph.runs):
                if run_idx < len(new_paragraph.runs):
                    new_run = new_paragraph.runs[run_idx]
                    new_run.bold = run.bold
                    new_run.italic = run.italic
                    new_run.underline = run.underline
                    if run.font.size:
                        new_run.font.size = run.font.size

        # Copy tables
        for table in source_doc.tables:
            copy_table(table, target_doc)

    target_doc.save(target_filename)
    return target_filename


def core_convert_to_pdf(filename: str, output_filename: str | None = None) -> str:
    """Convert a Word document to PDF format.

    Args:
        filename: Path to the Word document.
        output_filename: Optional path for the output PDF.

    Returns:
        The path to the created PDF.
    """
    if not output_filename:
        base_name, _ = os.path.splitext(filename)
        output_filename = f"{base_name}.pdf"
    elif not output_filename.lower().endswith(".pdf"):
        output_filename = f"{output_filename}.pdf"

    if not os.path.isabs(output_filename):
        output_filename = os.path.abspath(output_filename)

    output_dir = os.path.dirname(output_filename) or os.path.abspath(".")
    os.makedirs(output_dir, exist_ok=True)

    system = platform.system()

    if system == "Windows":
        try:
            from docx2pdf import convert
            convert(filename, output_filename)
            return output_filename
        except ImportError:
            raise RuntimeError("docx2pdf is required for PDF conversion on Windows. Please install it.")
        except Exception as e:
            raise RuntimeError(f"PDF conversion failed: {str(e)}")
            
    elif system in ["Linux", "Darwin"]:
        lo_commands = ["libreoffice", "soffice"] if system == "Linux" else ["soffice", "/Applications/LibreOffice.app/Contents/MacOS/soffice"]
        
        last_error = None
        for cmd_name in lo_commands:
            try:
                cmd = [cmd_name, "--headless", "--convert-to", "pdf", "--outdir", output_dir, filename]
                subprocess.run(cmd, capture_output=True, text=True, timeout=60, check=True)
                
                # Handle LibreOffice naming
                base_name = os.path.basename(filename)
                pdf_base_name = os.path.splitext(base_name)[0] + ".pdf"
                created_pdf = os.path.join(output_dir, pdf_base_name)
                
                if created_pdf != output_filename and os.path.exists(created_pdf):
                    shutil.move(created_pdf, output_filename)
                return output_filename
            except (subprocess.SubprocessError, FileNotFoundError) as e:
                last_error = e
                continue
        
        # Fallback to docx2pdf if available and LibreOffice failed
        try:
            from docx2pdf import convert
            convert(filename, output_filename)
            return output_filename
        except (ImportError, Exception):
            raise RuntimeError(f"PDF conversion failed. LibreOffice errors: {str(last_error)}. docx2pdf fallback also failed.")
    
    else:
        raise RuntimeError(f"Unsupported platform for PDF conversion: {system}")


def core_get_paragraph_text(doc: DocumentType, paragraph_index: int) -> dict[str, Any]:
    """Get text and metadata from a specific paragraph.

    Args:
        doc: The Word document object.
        paragraph_index: Index of the paragraph.

    Returns:
        Dictionary with paragraph text and metadata.
    """
    if not 0 <= paragraph_index < len(doc.paragraphs):
        raise IndexError(f"Paragraph index {paragraph_index} out of range.")

    paragraph: Paragraph = doc.paragraphs[paragraph_index]

    return {
        "index": paragraph_index,
        "text": paragraph.text,
        "style": paragraph.style.name if paragraph.style else "Normal",
        "is_heading": (
            paragraph.style.name.startswith("Heading") if paragraph.style else False
        ),
    }


def core_find_text(
    doc: DocumentType, text_to_find: str, match_case: bool = True, whole_word: bool = False
) -> dict[str, Any]:
    """Find all occurrences of specific text in a Word document.

    Args:
        doc: The Word document object.
        text_to_find: Text to search for.
        match_case: If True, performs a case-sensitive search.
        whole_word: If True, matches whole words only.

    Returns:
        A dictionary with search results.
    """
    if not text_to_find:
        raise ValueError("Search text cannot be empty")

    search_pattern = _create_search_pattern(text_to_find, match_case, whole_word)
    all_occurrences: list[dict[str, Any]] = []

    # Search in paragraphs
    for i, para in enumerate(doc.paragraphs):
        location = f"Paragraph {i}"
        all_occurrences.extend(_search_in_element(para, search_pattern, location))

    # Search in tables
    for t_idx, table in enumerate(doc.tables):
        for r_idx, row in enumerate(table.rows):
            for c_idx, cell in enumerate(row.cells):
                location = f"Table {t_idx}, Row {r_idx}, Cell {c_idx}"
                all_occurrences.extend(
                    _search_in_element(cell, search_pattern, location)
                )

    return {
        "query": text_to_find,
        "match_case": match_case,
        "whole_word": whole_word,
        "occurrences": all_occurrences,
        "total_count": len(all_occurrences),
    }


def _create_search_pattern(
    text_to_find: str, match_case: bool, whole_word: bool
) -> Pattern[str]:
    """Create a regex pattern for searching text."""
    pattern = re.escape(text_to_find)
    if whole_word:
        pattern = r"\b" + pattern + r"\b"

    flags = 0 if match_case else re.IGNORECASE
    return re.compile(pattern, flags)


def _search_in_element(
    element: Any, pattern: Pattern[str], location_prefix: str
) -> list[dict[str, Any]]:
    """Search for a pattern within a document element (paragraph or cell)."""
    occurrences = []
    for match in pattern.finditer(element.text):
        context = element.text[:100] + ("..." if len(element.text) > 100 else "")
        occurrences.append(
            {
                "location": location_prefix,
                "position": match.start(),
                "match": match.group(0),
                "context": context,
            }
        )
    return occurrences
