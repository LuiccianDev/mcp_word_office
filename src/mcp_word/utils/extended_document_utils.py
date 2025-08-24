"""
Extended document utilities for the Word Document Server.

This module provides advanced functions for interacting with Word documents,
including searching for text and retrieving specific content elements.
"""

import re
from re import Pattern
from typing import Any

from docx import Document
from docx.document import Document as DocumentType
from docx.text.paragraph import Paragraph

from mcp_word.validation.document_validators import validate_docx_file


@validate_docx_file("doc_path")
def get_paragraph_text(doc_path: str, paragraph_index: int) -> dict[str, Any]:
    """Get text from a specific paragraph in a Word document.

    Args:
        doc_path: Path to the Word document.
        paragraph_index: Index of the paragraph to extract (0-based).

    Returns:
        A dictionary with paragraph text and metadata, or an error dictionary.
    """
    try:
        doc: DocumentType = Document(doc_path)

        if not 0 <= paragraph_index < len(doc.paragraphs):
            return {
                "status": "error",
                "message": f"Invalid paragraph index: {paragraph_index}. "
                f"Document has {len(doc.paragraphs)} paragraphs.",
            }

        paragraph: Paragraph = doc.paragraphs[paragraph_index]

        return {
            "status": "success",
            "index": paragraph_index,
            "text": paragraph.text,
            "style": paragraph.style.name if paragraph.style else "Normal",
            "is_heading": (
                paragraph.style.name.startswith("Heading") if paragraph.style else False
            ),
        }
    except Exception as e:
        return {"status": "error", "message": f"Failed to get paragraph text: {e}"}


@validate_docx_file("doc_path")
def find_text(
    doc_path: str, text_to_find: str, match_case: bool = True, whole_word: bool = False
) -> dict[str, Any]:
    """Find all occurrences of specific text in a Word document.

    Args:
        doc_path: Path to the Word document.
        text_to_find: Text to search for.
        match_case: If True, performs a case-sensitive search.
        whole_word: If True, matches whole words only.

    Returns:
        A dictionary with search results, or an error dictionary.
    """
    if not text_to_find:
        return {"status": "error", "message": "Search text cannot be empty"}

    try:
        doc: DocumentType = Document(doc_path)
        search_pattern = _create_search_pattern(text_to_find, match_case, whole_word)
        all_occurrences: list[dict[str, Any]] = []

        # Search in paragraphs
        for i, para in enumerate(doc.paragraphs):
            location = f"Paragraph {i}"
            all_occurrences.extend(_search_in_element(para, search_pattern, location))

        # Search in tables
        for t_idx, table in enumerate(doc.tables):
            for r_idx, row in enumerate(table.rows):
                for c_idx, cell in enumerate(row.cells):
                    location = f"Table {t_idx}, Row {r_idx}, Cell {c_idx}"
                    all_occurrences.extend(
                        _search_in_element(cell, search_pattern, location)
                    )

        return {
            "status": "success",
            "query": text_to_find,
            "match_case": match_case,
            "whole_word": whole_word,
            "occurrences": all_occurrences,
            "total_count": len(all_occurrences),
        }
    except Exception as e:
        return {"status": "error", "message": f"Failed to search for text: {e}"}


def _create_search_pattern(
    text_to_find: str, match_case: bool, whole_word: bool
) -> Pattern[str]:
    """Create a regex pattern for searching text."""
    pattern = re.escape(text_to_find)
    if whole_word:
        pattern = r"\b" + pattern + r"\b"

    flags = 0 if match_case else re.IGNORECASE
    return re.compile(pattern, flags)


def _search_in_element(
    element: Any, pattern: Pattern[str], location_prefix: str
) -> list[dict[str, Any]]:
    """Search for a pattern within a document element (paragraph or cell)."""
    occurrences = []
    for match in pattern.finditer(element.text):
        context = element.text[:100] + ("..." if len(element.text) > 100 else "")
        occurrences.append(
            {
                "location": location_prefix,
                "position": match.start(),
                "match": match.group(0),
                "context": context,
            }
        )
    return occurrences
