"""Tests for the document_utils.py module."""
import os
import pytest
from pathlib import Path
from datetime import datetime
from unittest.mock import patch, MagicMock

# Import the module to test
from mcp_word_server.utils.document_utils import get_document_properties

# Test data directory path
TEST_DATA_DIR = Path(__file__).parent / "test_data"
SAMPLE_DOCX = TEST_DATA_DIR / "test_document.docx"

@pytest.fixture(scope="module", autouse=True)
def setup_test_files():
    """Setup test files before tests and clean up afterward."""
    # Create test_data directory if it doesn't exist
    TEST_DATA_DIR.mkdir(exist_ok=True)
    
    # Create a sample .docx file with known properties
    from docx import Document
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    
    doc = Document()
    
    # Add document properties
    core_props = doc.core_properties
    core_props.title = "Test Document"
    core_props.author = "Test Author"
    core_props.subject = "Test Subject"
    core_props.keywords = "test, document, pytest"
    
    # Add some content
    doc.add_heading("Test Document", level=1)
    doc.add_paragraph("This is a test document for unit testing.")
    
    # Add a table
    table = doc.add_table(rows=1, cols=2)
    row = table.rows[0].cells
    row[0].text = "Header 1"
    row[1].text = "Header 2"
    
    # Save the document
    doc.save(SAMPLE_DOCX)
    
    yield  # This is where the testing happens
    
    # Cleanup
    if SAMPLE_DOCX.exists():
        SAMPLE_DOCX.unlink()
    if TEST_DATA_DIR.exists() and not any(TEST_DATA_DIR.iterdir()):
        TEST_DATA_DIR.rmdir()

def test_get_document_properties_success():
    """Test getting document properties with a valid document."""
    # Execute
    result = get_document_properties(str(SAMPLE_DOCX))
    
    # Assert
    assert isinstance(result, dict)
    assert "error" not in result
    assert result["title"] == "Test Document"
    assert result["author"] == "Test Author"
    assert result["subject"] == "Test Subject"
    assert "test" in result["keywords"]
    assert result["table_count"] == 1
    assert result["paragraph_count"] > 0
    assert result["word_count"] > 0

def test_get_document_properties_nonexistent_file():
    """Test getting properties with a non-existent file."""
    # Execute
    result = get_document_properties("nonexistent.docx")
    
    # Assert
    assert isinstance(result, dict)
    assert "error" in result
    assert "does not exist" in result["error"]

def test_get_document_properties_invalid_file():
    """Test getting properties with an invalid file."""
    # Create an invalid file
    invalid_file = TEST_DATA_DIR / "invalid.docx"
    invalid_file.write_text("This is not a valid docx")
    
    try:
        # Execute
        result = get_document_properties(str(invalid_file))
        
        # Assert
        assert isinstance(result, dict)
        assert "error" in result
        assert "not a valid Word document" in result["error"]
    finally:
        # Cleanup
        if invalid_file.exists():
            invalid_file.unlink()

@patch('mcp_word_server.utils.document_utils.Document')
def test_get_document_properties_exception_handling(mock_document):
    """Test exception handling in get_document_properties."""
    # Setup mock to raise an exception
    mock_document.side_effect = Exception("Test exception")
    
    # Execute
    result = get_document_properties(str(SAMPLE_DOCX))
    
    # Assert
    assert isinstance(result, dict)
    assert "error" in result
    assert "Failed to get document properties" in result["error"]

def test_get_document_properties_empty_properties():
    """Test with a document that has no properties set."""
    # Create a minimal document
    from docx import Document
    empty_doc = TEST_DATA_DIR / "empty.docx"
    
    try:
        doc = Document()
        # Clear default properties
        doc.core_properties.title = ""
        doc.core_properties.author = ""
        doc.save(empty_doc)
    
        # Execute
        result = get_document_properties(str(empty_doc))
        
        # Assert
        assert isinstance(result, dict)
        assert "error" not in result
        assert result["title"] == ""
        assert result["author"] == ""
        assert result["table_count"] == 0
    finally:
        # Cleanup
        if empty_doc.exists():
            empty_doc.unlink()
