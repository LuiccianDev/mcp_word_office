"""Tests for the header_footer_tools module."""

from collections.abc import Generator
from pathlib import Path
import pytest
from docx import Document

from mcp_word.tools import header_footer_tools


@pytest.fixture
def temp_docx_file(tmp_path: Path) -> Generator[str, None, None]:
    """Create a temporary .docx file for testing."""
    file_path = tmp_path / "test_document.docx"
    doc = Document()
    doc.add_paragraph("This is a test document.")
    doc.save(file_path)
    yield str(file_path)


@pytest.mark.asyncio
async def test_add_header(temp_docx_file: str) -> None:
    """Test adding a header."""
    result = await header_footer_tools.add_header(temp_docx_file, "Test Header Text")
    assert result["status"] == "success"
    
    # Verify
    doc = Document(temp_docx_file)
    assert doc.sections[0].header.paragraphs[0].text == "Test Header Text"


@pytest.mark.asyncio
async def test_add_footer(temp_docx_file: str) -> None:
    """Test adding a footer."""
    result = await header_footer_tools.add_footer(temp_docx_file, "Test Footer Text")
    assert result["status"] == "success"
    
    # Verify
    doc = Document(temp_docx_file)
    assert doc.sections[0].footer.paragraphs[0].text == "Test Footer Text"


@pytest.mark.asyncio
async def test_add_header_invalid_section(temp_docx_file: str) -> None:
    """Test adding a header to an invalid section."""
    result = await header_footer_tools.add_header(temp_docx_file, "Test", section_index=5)
    assert result["status"] == "error"
    assert "out of bounds" in result["message"]
