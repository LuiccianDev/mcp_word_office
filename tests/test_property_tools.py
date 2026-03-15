"""Tests for the property_tools module."""

from collections.abc import Generator
from pathlib import Path
import pytest
from docx import Document

from mcp_word.tools import property_tools


@pytest.fixture
def temp_docx_file(tmp_path: Path) -> Generator[str, None, None]:
    """Create a temporary .docx file for testing."""
    file_path = tmp_path / "test_document.docx"
    doc = Document()
    doc.add_paragraph("Test.")
    doc.save(file_path)
    yield str(file_path)


@pytest.mark.asyncio
async def test_set_and_get_core_properties(temp_docx_file: str) -> None:
    """Test setting and getting core document properties."""
    set_result = await property_tools.set_core_properties(
        temp_docx_file,
        author="Test Author",
        title="Test Title"
    )
    assert set_result["status"] == "success"

    get_result = await property_tools.get_core_properties(temp_docx_file)
    assert get_result["status"] == "success"
    assert get_result["properties"]["author"] == "Test Author"
    assert get_result["properties"]["title"] == "Test Title"


@pytest.mark.asyncio
async def test_set_empty_properties(temp_docx_file: str) -> None:
    """Test that setting no properties returns an error."""
    result = await property_tools.set_core_properties(temp_docx_file)
    assert result["status"] == "error"


@pytest.mark.asyncio
async def test_set_page_layout(temp_docx_file: str) -> None:
    """Test setting the page layout."""
    result = await property_tools.set_page_layout(temp_docx_file, 0, "landscape")
    assert result["status"] == "success"

    doc = Document(temp_docx_file)
    assert doc.sections[0].orientation == 1  # WD_ORIENTATION.LANDSCAPE
    assert int(doc.sections[0].page_width.inches) == 11
