"""Tests for the content_tools module."""

from collections.abc import Generator
from pathlib import Path
from unittest.mock import MagicMock, patch

import pytest
from docx import Document
from PIL import Image

# Assuming content_tools.py is in src/word_mcp/tools
from mcp_word.tools import content_tools


@pytest.fixture  # type: ignore[misc]
def temp_docx_file(tmp_path: Path) -> Generator[str, None, None]:
    """Create a temporary .docx file for testing."""
    file_path = tmp_path / "test_document.docx"
    doc = Document()
    doc.add_paragraph("This is a test document.")
    doc.save(file_path)
    yield str(file_path)


@pytest.mark.asyncio  # type: ignore[misc]
async def test_add_heading(temp_docx_file: str) -> None:
    """Test adding a heading to a document."""
    result = await content_tools.add_heading(temp_docx_file, "Test Heading", level=1)
    assert result["status"] == "success"
    assert "Heading" in result["message"]


@pytest.mark.asyncio  # type: ignore[misc]
async def test_add_paragraph(temp_docx_file: str) -> None:
    """Test adding a paragraph to a document."""
    result = await content_tools.add_paragraph(
        temp_docx_file, "This is a new paragraph."
    )
    assert result["status"] == "success"
    assert "Paragraph" in result["message"]


@pytest.mark.asyncio  # type: ignore[misc]
async def test_add_table(temp_docx_file: str) -> None:
    """Test adding a table to a document."""
    result = await content_tools.add_table(temp_docx_file, rows=2, cols=3)
    assert result["status"] == "success"
    assert "Table" in result["message"]
    assert "2 rows" in result["message"]
    assert "3 columns" in result["message"]


@pytest.mark.asyncio  # type: ignore[misc]
async def test_add_picture_with_patch(temp_docx_file: str, tmp_path: Path) -> None:
    """Test adding a picture using patch."""

    image_path = tmp_path / "test_image.png"
    image = Image.new("RGB", (100, 100), color="red")
    image.save(image_path)

    with patch("mcp_word.tools.content_tools.core_add_picture") as mock_core:
        with patch("mcp_word.tools.content_tools.document_context") as mock_ctx:
            mock_doc = MagicMock()
            mock_ctx.return_value.__enter__.return_value = mock_doc
            mock_core.return_value = {
                "image_name": "test_image.png",
                "image_path": str(image_path),
                "width_inches": None
            }
            
            result = await content_tools.add_picture(temp_docx_file, str(image_path))
            
            # Debug print to see what we actually got
            print(f"DEBUG mock_core.call_args: {mock_core.call_args}")
            
            mock_core.assert_called_once()
            args, kwargs = mock_core.call_args
            assert args[0] == mock_doc
            assert Path(args[1]).resolve() == Path(image_path).resolve()
            # Handle both positional and keyword width
            actual_width = args[2] if len(args) > 2 else kwargs.get('width')
            assert actual_width is None
            assert result["status"] == "success"


@pytest.mark.asyncio  # type: ignore[misc]
async def test_add_page_break(temp_docx_file: str) -> None:
    """Test adding a page break to a document."""
    result = await content_tools.add_page_break(temp_docx_file)
    assert result["status"] == "success"
    assert "Page break" in result["message"]


@pytest.mark.asyncio  # type: ignore[misc]
async def test_add_table_of_contents(temp_docx_file: str) -> None:
    """Test adding a table of contents to a document."""
    await content_tools.add_heading(temp_docx_file, "TOC Heading", 1)
    result = await content_tools.add_table_of_contents(temp_docx_file)
    assert result["status"] == "success"
    assert "Table of Contents" in result["message"] or "TOC" in result["message"]


@pytest.mark.asyncio  # type: ignore[misc]
async def test_delete_paragraph(temp_docx_file: str) -> None:
    """Test deleting a paragraph from a document."""
    await content_tools.add_paragraph(temp_docx_file, "Paragraph to be deleted.")
    result = await content_tools.delete_paragraph(
        temp_docx_file, 1
    )  # Delete the second paragraph
    assert result["status"] == "success"
    assert "Paragraph" in result["message"] and "deleted" in result["message"]


@pytest.mark.asyncio  # type: ignore[misc]
async def test_search_and_replace(temp_docx_file: str) -> None:
    """Test searching and replacing text in a document."""
    result = await content_tools.search_and_replace(temp_docx_file, "test", "sample")
    assert result["status"] == "success"
    assert "Replaced" in result["message"] or "occurrences" in result["message"]
