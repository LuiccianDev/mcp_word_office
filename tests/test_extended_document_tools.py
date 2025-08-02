"""Tests for the extended_document_tools module."""

from pathlib import Path
from typing import Generator
from unittest.mock import Mock, patch

import pytest
from docx import Document

# Assuming extended_document_tools.py is in src/word_mcp/tools
from word_mcp.tools import extended_document_tools


@pytest.fixture  # type: ignore[misc]
def temp_docx_file(tmp_path: Path) -> Generator[str, None, None]:
    """Create a temporary .docx file for testing."""
    file_path = tmp_path / "test_document.docx"
    doc = Document()
    doc.add_paragraph("This is the first paragraph.")
    doc.add_paragraph("This is the second paragraph with the word Test.")
    doc.save(file_path)
    yield str(file_path)


@pytest.mark.asyncio  # type: ignore[misc]
async def test_get_paragraph_text_from_document(temp_docx_file: str) -> None:
    """Test getting text from a specific paragraph."""
    result = await extended_document_tools.get_paragraph_text_from_document(
        temp_docx_file, 1
    )
    assert result["status"] == "success"
    assert "second paragraph" in result["text"]


@pytest.mark.asyncio  # type: ignore[misc]
async def test_find_text_in_document(temp_docx_file: str) -> None:
    """Test finding text in a document."""
    result = await extended_document_tools.find_text_in_document(temp_docx_file, "Test")
    assert result["status"] == "success"


@pytest.mark.asyncio  # type: ignore[misc]
@patch("platform.system", return_value="Windows")
@patch("docx2pdf.convert")
async def test_convert_to_pdf_windows(
    mock_convert: Mock, mock_system: Mock, temp_docx_file: str, tmp_path: Path
) -> None:
    """Test converting a document to PDF on Windows."""
    output_file = tmp_path / "output.pdf"
    result = await extended_document_tools.convert_to_pdf(
        temp_docx_file, str(output_file)
    )
    mock_convert.assert_called_once_with(temp_docx_file, str(output_file))
    assert result["success"]
    assert "successfully converted" in result["message"]
