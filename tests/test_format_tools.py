"""Tests for the format_tools module."""

from pathlib import Path
from typing import Generator

import pytest
from docx import Document

# Assuming format_tools.py is in src/word_mcp/tools
from mcp_word.tools import format_tools


@pytest.fixture  # type: ignore[misc]
def temp_docx_file_with_content(tmp_path: Path) -> Generator[str, None, None]:
    """Create a temporary .docx file with content for testing."""
    file_path = tmp_path / "test_document.docx"
    doc = Document()
    doc.add_paragraph("This is a test paragraph for formatting.")
    doc.add_table(rows=2, cols=2, style="Table Grid")
    doc.save(file_path)
    yield str(file_path)


@pytest.mark.asyncio  # type: ignore[misc]
async def test_format_text(temp_docx_file_with_content: str) -> None:
    """Test formatting a range of text in a paragraph."""
    result = await format_tools.format_text(
        temp_docx_file_with_content, 0, 5, 7, bold=True, italic=True
    )
    assert result["status"] == "success"


@pytest.mark.asyncio  # type: ignore[misc]
async def test_format_table(temp_docx_file_with_content: str) -> None:
    """Test formatting a table."""
    result = await format_tools.format_table(
        temp_docx_file_with_content, 0, has_header_row=True, border_style="single"
    )
    assert result["status"] == "success"
